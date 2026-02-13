#!/usr/bin/env python3
"""
revise_manuscript.py

Applies V1/V2 dual-model framework revisions to:
  - BMC_Public_Health_Manuscript.docx
  - BMC_Cover_Letter.docx

Uses python-docx to perform surgical paragraph-level edits.
Input files are not modified; revised copies are saved with _revised suffix.
"""

import copy
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


# ─── paths ────────────────────────────────────────────────────────────────────
INPUT_DIR = Path("/home/claude/HIV_Prevention_PWID/manuscript_prep")
OUTPUT_DIR = INPUT_DIR  # same directory, _revised suffix

MANUSCRIPT_IN = INPUT_DIR / "BMC_Public_Health_Manuscript.docx"
MANUSCRIPT_OUT = OUTPUT_DIR / "BMC_Public_Health_Manuscript_revised.docx"

COVER_IN = INPUT_DIR / "BMC_Cover_Letter.docx"
COVER_OUT = OUTPUT_DIR / "BMC_Cover_Letter_revised.docx"


# ─── helpers ──────────────────────────────────────────────────────────────────

def find_paragraph(doc, prefix, start=0):
    """Return the index of the first paragraph whose text starts with *prefix*."""
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if p.text.strip().startswith(prefix):
            return i
    raise ValueError(f"Paragraph starting with {prefix!r} not found (start={start})")


def replace_paragraph_text(paragraph, new_text):
    """Replace all text in a paragraph while preserving the style of the first run."""
    # Preserve the paragraph's style
    style = paragraph.style
    # Capture formatting from the first run (if any)
    first_run_font = None
    if paragraph.runs:
        r = paragraph.runs[0]
        first_run_font = {
            'bold': r.bold,
            'italic': r.italic,
            'size': r.font.size,
            'name': r.font.name,
        }

    # Clear all runs
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)

    # Also clear any remaining text nodes directly in the paragraph XML
    p_elem = paragraph._element
    for child in list(p_elem):
        if child.tag == qn('w:r'):
            p_elem.remove(child)

    # Add new run with the new text
    new_run = paragraph.add_run(new_text)
    if first_run_font:
        if first_run_font['bold'] is not None:
            new_run.bold = first_run_font['bold']
        if first_run_font['italic'] is not None:
            new_run.italic = first_run_font['italic']
        if first_run_font['size'] is not None:
            new_run.font.size = first_run_font['size']
        if first_run_font['name'] is not None:
            new_run.font.name = first_run_font['name']

    paragraph.style = style


def insert_paragraph_after(doc, ref_paragraph, text, style=None):
    """Insert a new paragraph immediately after *ref_paragraph* in the document body.

    Returns the newly inserted Paragraph object (from doc.paragraphs refresh).
    """
    from lxml import etree
    from docx.oxml import OxmlElement

    # Create a new <w:p> element
    new_p = OxmlElement('w:p')

    # Copy paragraph properties (pPr) from reference if present
    ref_pPr = ref_paragraph._element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_p.append(copy.deepcopy(ref_pPr))

    # Create run with text
    new_r = OxmlElement('w:r')

    # Copy run properties from reference first run
    if ref_paragraph.runs:
        ref_rPr = ref_paragraph.runs[0]._element.find(qn('w:rPr'))
        if ref_rPr is not None:
            new_r.append(copy.deepcopy(ref_rPr))

    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)

    # Insert after the reference paragraph in the XML tree
    ref_paragraph._element.addnext(new_p)

    # Find and return the new paragraph via doc.paragraphs
    # (This gives us a properly bound Paragraph object)
    for p in doc.paragraphs:
        if p._element is new_p:
            return p

    # Fallback: return a lightweight wrapper (shouldn't happen)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, doc.element.body)


# ─── MANUSCRIPT EDITS ─────────────────────────────────────────────────────────

def revise_manuscript():
    doc = Document(str(MANUSCRIPT_IN))

    # NOTE: We always use doc.paragraphs[idx] (fresh property access) because
    # inserting paragraphs changes the list.  Never cache doc.paragraphs.

    # ── P9: Abstract Methods ──────────────────────────────────────────────
    idx = find_paragraph(doc, "Methods: We developed a Monte Carlo")
    replace_paragraph_text(doc.paragraphs[idx], (
        "Methods: We developed a Monte Carlo simulation modeling an 8-step "
        "LAI-PrEP cascade for PWID under current U.S. policy conditions "
        "(n = 100,000 per scenario). We decomposed barriers into three layers: "
        "pathogen biology, HIV testing gaps, and architectural barriers. We "
        "compared PWID outcomes to men who have sex with men (MSM) receiving "
        "identical pharmacological interventions and modeled stochastic "
        "avoidance failure using network density dynamics. We present two "
        "model versions: a baseline stochastic framework (V1) and an enhanced "
        "model (V2) incorporating a multiplicative methamphetamine × housing "
        "interaction derived from Hood et al. (2018), which captures synergistic "
        "structural vulnerability where co-occurring methamphetamine use and "
        "unstable housing accelerate network densification beyond the sum of "
        "individual effects."
    ))

    # ── P10: Abstract Results ─────────────────────────────────────────────
    idx = find_paragraph(doc, "Results: Under current policy, PWID achieved")
    replace_paragraph_text(doc.paragraphs[idx], (
        "Results: Under current policy, PWID achieved P(R\u2080 = 0) = 0.003% "
        "(95% CI: 0.000\u20130.006%) compared to 16.3% for MSM\u2014a 5,434-fold "
        "disparity. Architectural barriers accounted for 93.1% of cascade "
        "failure. Criminalization was the single largest contributor (38.4%). "
        "The V1 stochastic avoidance model predicted 73.0% probability of "
        "major outbreak within 5 years (median: 3.0 years). The V2 model, "
        "incorporating meth \u00d7 housing interaction, reduced national 5-year "
        "risk to 68.4% but increased Pacific Northwest risk to 92.4% "
        "(median: 1.0 year), demonstrating that structural intersections "
        "can accelerate regional outbreak timelines."
    ))

    # ── P11: Abstract Conclusions ─────────────────────────────────────────
    idx = find_paragraph(doc, "Conclusions: Current HIV prevention for PWID")
    replace_paragraph_text(doc.paragraphs[idx], (
        "Conclusions: Current HIV prevention for PWID relies on probability "
        "rather than intervention. Structural barriers create conditions where "
        "effective prevention is mathematically infeasible regardless of drug "
        "efficacy. The 5,434-fold disparity represents policy choices, not "
        "epidemic inevitability. Together, V1 and V2 bracket plausible outbreak "
        "risk at 68\u201373% nationally within 5 years, with structural interactions "
        "accelerating regional collapse to as little as 1 year in high-burden "
        "settings."
    ))

    # ── P28: Stochastic avoidance model body ──────────────────────────────
    idx = find_paragraph(doc, "Network density was modeled as a function")
    replace_paragraph_text(doc.paragraphs[idx], (
        "Network density was modeled as a function of exogenous contextual "
        "drivers, including stimulant-associated network expansion documented "
        "in outbreak investigations, housing instability (68.5%) [14, 15], "
        "and sex-work bridging. These drivers are treated as structural modifiers "
        "of network connectivity rather than individual-level behavioral "
        "determinants. Annual outbreak probability increases exponentially "
        "above the critical density threshold, modulated by syringe service "
        "program and opioid agonist therapy coverage [16\u201318]."
    ))

    # Insert V1/V2 explanation paragraph after P28
    new_para = insert_paragraph_after(doc, doc.paragraphs[idx], (
        "We present two model versions. The baseline model (V1) treats "
        "methamphetamine and housing instability effects as additive "
        "contributors to network density. An enhanced model (V2) incorporates "
        "a multiplicative methamphetamine \u00d7 housing interaction term "
        "(coefficient = 0.8; range 0.3\u20131.5) derived from Hood et al. (2018), "
        "who demonstrated that the joint effect of methamphetamine use and "
        "unstable housing on viral suppression (42% vs. 76% baseline) exceeds "
        "the sum of individual effects by approximately 1.5-fold [21]. V2 "
        "captures synergistic structural vulnerability: when both "
        "methamphetamine prevalence and housing instability are elevated, "
        "network density increases beyond what additive models predict. When "
        "the interaction coefficient is set to zero, V2 reproduces V1 outputs, "
        "confirming backward compatibility. Full model specification is "
        "provided in Additional file 2."
    ))

    # ── P30: Monte Carlo simulation body ──────────────────────────────────
    # Re-find because we inserted a paragraph
    idx = find_paragraph(doc, "Monte Carlo simulations (n = 100,000")
    replace_paragraph_text(doc.paragraphs[idx], (
        "Monte Carlo simulations (n = 100,000 individuals per scenario for "
        "cascade analysis; n = 2,000 simulations each for V1 and V2 outbreak "
        "modeling) estimated the probability of achieving sustained protection "
        "under current and counterfactual policy conditions. Sensitivity analyses "
        "including tornado diagrams, policy scenario comparisons, and "
        "probabilistic sensitivity analysis assessed robustness across parameter "
        "uncertainty. Detailed methods are provided in Additional file 1."
    ))

    # ── P50: Stochastic avoidance failure body ────────────────────────────
    idx = find_paragraph(doc, "The stochastic avoidance model predicted 73.8%")
    replace_paragraph_text(doc.paragraphs[idx], (
        "The V1 stochastic avoidance model predicted 73.0% probability of "
        "major outbreak within 5 years under current conditions (Fig. 4). "
        "Median time to outbreak was 3.0 years, with cumulative probability "
        "reaching 93.4% by 10 years. Regional variation was substantial: "
        "Pacific Northwest showed 86.8% 5-year probability (median 2.0 years), "
        "consistent with elevated methamphetamine prevalence and housing "
        "instability in this region."
    ))

    # Insert V2 results paragraph
    v2_para = insert_paragraph_after(doc, doc.paragraphs[idx], (
        "The V2 model, incorporating multiplicative meth \u00d7 housing interaction, "
        "demonstrated divergent regional effects (Fig. 6). Nationally, V2 "
        "reduced the 5-year outbreak probability to 68.4% (\u22124.6 percentage "
        "points from V1), as the interaction term combined with the exponential "
        "threshold creates a different sensitivity landscape at the national "
        "level. However, in the Pacific Northwest\u2014where both methamphetamine "
        "prevalence (35%) and housing instability are simultaneously elevated\u2014"
        "V2 raised the 5-year outbreak probability to 92.4% (+5.7 percentage "
        "points) and halved the median time to outbreak from 2.0 to 1.0 years."
    ))

    # Insert bracketing paragraph after the V2 paragraph
    insert_paragraph_after(doc, v2_para, (
        "Together, V1 and V2 bracket the plausible range of national 5-year "
        "outbreak risk at 68\u201373%. Sensitivity analysis identified the meth "
        "\u00d7 housing interaction coefficient as the 6th most influential "
        "parameter (11.2 percentage point range), ranking between housing "
        "instability rate and methamphetamine growth rate in the tornado "
        "diagram (Fig. 7). Probabilistic sensitivity analysis yielded a V2 "
        "mean 5-year outbreak probability of 64.7% (90% CI: 32.0\u201398.0%). "
        "Full regional analysis, tornado diagrams, and scenario comparisons "
        "are provided in Additional file 2."
    ))

    # ── P54: Additional analyses sentence — add V2 supplementary figs ─────
    idx = find_paragraph(doc, "Additional analyses examining regional heterogeneity")
    replace_paragraph_text(doc.paragraphs[idx], (
        "Additional analyses examining regional heterogeneity, outbreak "
        "probability, and sensitivity of stochastic avoidance dynamics are "
        "provided in the additional files and associated code repository. "
        "Additional analyses examining V1 versus V2 model comparison, "
        "including network density trajectories (Fig. S5), outbreak forecast "
        "comparison (Fig. S6), V2 tornado diagram (Fig. S7), and V2 scenario "
        "comparison (Fig. S8), are provided in the additional files."
    ))

    # ── P60: Discussion — three findings, revise second ───────────────────
    idx = find_paragraph(doc, "Three findings warrant particular attention")
    original_text = doc.paragraphs[idx].text
    # Replace the second finding sentence about 73.8%
    old_second = (
        "Second, the 73.8% five-year outbreak probability represents "
        "predictable system failure, not epidemic randomness; current "
        "prevention relies on stochastic avoidance that network density "
        "trends are actively eroding."
    )
    new_second = (
        "Second, the V1 model\u2019s 73.0% and V2 model\u2019s 68.4% five-year "
        "national outbreak probabilities represent predictable system failure, "
        "not epidemic randomness; current prevention relies on stochastic "
        "avoidance that network density trends are actively eroding. The V2 "
        "model\u2019s divergent regional behavior\u2014lowering national risk while "
        "accelerating Pacific Northwest outbreak timelines to a median of "
        "1.0 year\u2014demonstrates that structural intersections between "
        "methamphetamine use and housing instability can transform gradual "
        "erosion into acute regional collapse."
    )
    new_text = original_text.replace(old_second, new_second)
    replace_paragraph_text(doc.paragraphs[idx], new_text)

    # ── P61: Discussion — add regional divergence sentence ────────────────
    idx = find_paragraph(doc, "Exogenous increases in network density")
    original_text = doc.paragraphs[idx].text
    replace_paragraph_text(doc.paragraphs[idx], (
        original_text + " This regional divergence underscores a critical "
        "policy implication: interventions designed around national-level "
        "risk estimates may systematically underestimate outbreak urgency "
        "in communities where multiple structural vulnerabilities intersect."
    ))

    # ── P64: Limitations ──────────────────────────────────────────────────
    idx = find_paragraph(doc, "Our barrier parameterization relies on")
    replace_paragraph_text(doc.paragraphs[idx], (
        "Our barrier parameterization relies on heterogeneous literature "
        "sources with varying methodological quality. The stochastic avoidance "
        "model simplifies complex network dynamics into aggregate density "
        "measures. We assume barrier effects are multiplicatively independent, "
        "which may underestimate synergistic interactions\u2014though the V2 model "
        "partially addresses this limitation by incorporating an empirically-derived "
        "meth \u00d7 housing interaction term from Hood et al. (2018). The "
        "interaction coefficient (0.8; range 0.3\u20131.5) was estimated from King "
        "County surveillance data and may not generalize to all regions. "
        "Regional variation is modeled at aggregate level rather than "
        "county-specific dynamics."
    ))

    # ── P75: Data Availability ────────────────────────────────────────────
    idx = find_paragraph(doc, "All code and data supporting the findings")
    replace_paragraph_text(doc.paragraphs[idx], (
        "All code and data supporting the findings of this study are publicly "
        "available at: https://github.com/Nyx-Dynamics/HIV_Prevention_PWID. "
        "Key files include: architectural_barrier_model.py (prevention cascade "
        "simulation), cascade_sensitivity_analysis.py (probabilistic sensitivity "
        "analysis), stochastic_avoidance_enhanced.py (V1 stochastic avoidance "
        "and outbreak modeling), stochastic_avoidance_v2.py (V2 model with "
        "meth \u00d7 housing interaction), hood_parameter_comparison.py (Hood et al. "
        "calibration comparison), and config/parameters.json (model parameters "
        "with literature sources)."
    ))

    # ── P111: Fig. 4 legend — update 73.8% → 73.0%, 92.7% → 93.4% ──────
    idx = find_paragraph(doc, "Fig. 4 Stochastic avoidance failure prediction")
    replace_paragraph_text(doc.paragraphs[idx], (
        "Fig. 4 Stochastic avoidance failure prediction. (a) Cumulative "
        "outbreak probability over time, reaching 73.0% at 5 years and "
        "93.4% at 10 years. (b) Time-to-outbreak distribution showing median "
        "of 3.0 years (interquartile range 1\u20136 years). Based on n = 2,000 "
        "Monte Carlo simulations under current policy conditions"
    ))

    # Insert Fig. 6 legend after Fig. 4 legend
    # Need to skip past the blank paragraph after Fig. 4
    # The blank P112 follows, then P113 is Fig. 5 legend
    fig6_legend = insert_paragraph_after(doc, doc.paragraphs[idx + 1], (
        "Fig. 6 V1 versus V2 stochastic avoidance model comparison. "
        "(a) Cumulative outbreak probability: V1 (additive model, blue) "
        "versus V2 (meth \u00d7 housing interaction, red) with 90% simulation "
        "intervals. V1 predicts 73.0% national 5-year risk; V2 predicts "
        "68.4%. (b) Time-to-outbreak distributions showing overlapping V1 "
        "and V2 histograms. Both models show median 3.0-year national "
        "outbreak timing. Based on n = 2,000 Monte Carlo simulations per model."
    ))

    # Insert blank paragraph after Fig. 6 legend
    blank_after_fig6 = insert_paragraph_after(doc, fig6_legend, "")

    # Insert Fig. 7 legend after blank
    fig7_legend = insert_paragraph_after(doc, blank_after_fig6, (
        "Fig. 7 V2 tornado sensitivity analysis. One-way sensitivity of "
        "5-year outbreak probability to parameter variation. The meth \u00d7 "
        "housing interaction coefficient (green bars, labeled \u201cV2 NEW\u201d) "
        "ranks 6th, with 11.2 percentage point range. Baseline annual "
        "outbreak probability dominates (52.7pp range). Colors distinguish "
        "original parameters (blue/red) from the V2 interaction parameter "
        "(green)."
    ))

    # Insert blank paragraph after Fig. 7 legend
    insert_paragraph_after(doc, fig7_legend, "")

    # ── Table 2: Expand with V2 columns ───────────────────────────────────
    # Table 1 (index 1 in doc.tables) is Table 2 in the manuscript
    table2 = doc.tables[1]

    # New header row
    new_headers = [
        "Region", "V1 P(5yr)", "V2 P(5yr)",
        "V1 P(10yr)", "V2 P(10yr)", "V1 Median", "V2 Median"
    ]

    # New data rows
    new_data = [
        ["Pacific Northwest", "86.8%", "92.4%", "97.7%", "99.3%", "2.0", "1.0"],
        ["Appalachia", "78.4%", "—", "94.9%", "—", "2.0", "—"],
        ["Northeast Urban", "78.3%", "—", "94.2%", "—", "2.0", "—"],
        ["National Average", "73.0%", "68.4%", "93.4%", "90.7%", "3.0", "3.0"],
    ]

    # Clear existing table and rebuild
    # Remove existing rows (keep first for structure reference)
    tbl_elem = table2._tbl

    # Get the table grid and properties
    tblPr = tbl_elem.find(qn('w:tblPr'))
    tblGrid = tbl_elem.find(qn('w:tblGrid'))

    # Copy formatting from existing cells for reference
    ref_row = table2.rows[0]
    ref_cell_format = []
    for cell in ref_row.cells:
        tc = cell._element
        tcPr = tc.find(qn('w:tcPr'))
        ref_cell_format.append(copy.deepcopy(tcPr) if tcPr is not None else None)

    # Reference a data row for cell formatting
    data_ref_row = table2.rows[1]
    data_cell_format = []
    for cell in data_ref_row.cells:
        tc = cell._element
        tcPr = tc.find(qn('w:tcPr'))
        data_cell_format.append(copy.deepcopy(tcPr) if tcPr is not None else None)

    # Get run formatting from header and data cells
    header_run_props = None
    if ref_row.cells[0].paragraphs[0].runs:
        r = ref_row.cells[0].paragraphs[0].runs[0]
        header_run_props = copy.deepcopy(r._element.find(qn('w:rPr')))

    data_run_props = None
    if data_ref_row.cells[0].paragraphs[0].runs:
        r = data_ref_row.cells[0].paragraphs[0].runs[0]
        data_run_props = copy.deepcopy(r._element.find(qn('w:rPr')))

    # Remove all existing rows
    for row in list(tbl_elem.findall(qn('w:tr'))):
        tbl_elem.remove(row)

    # Update tblGrid to have 7 columns
    if tblGrid is not None:
        for gc in list(tblGrid.findall(qn('w:gridCol'))):
            tblGrid.remove(gc)
        for _ in range(7):
            gridCol = copy.deepcopy(tblGrid.makeelement(qn('w:gridCol'), {}))
            gridCol.set(qn('w:w'), '1300')
            tblGrid.append(gridCol)

    def add_table_row(tbl_elem, cells_text, is_header=False):
        """Add a row to the table XML."""
        tr = tbl_elem.makeelement(qn('w:tr'), {})
        for i, text in enumerate(cells_text):
            tc = tr.makeelement(qn('w:tc'), {})
            # Add cell properties for width
            tcPr = tc.makeelement(qn('w:tcPr'), {})
            tcW = tcPr.makeelement(qn('w:tcW'), {})
            tcW.set(qn('w:w'), '1300')
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            tc.append(tcPr)

            p = tc.makeelement(qn('w:p'), {})
            r = p.makeelement(qn('w:r'), {})

            # Apply run properties
            rPr_source = header_run_props if is_header else data_run_props
            if rPr_source is not None:
                r.append(copy.deepcopy(rPr_source))

            t = r.makeelement(qn('w:t'), {})
            t.text = text
            t.set(qn('xml:space'), 'preserve')
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl_elem.append(tr)

    # Add header row
    add_table_row(tbl_elem, new_headers, is_header=True)

    # Add data rows
    for row_data in new_data:
        add_table_row(tbl_elem, row_data, is_header=False)

    # ── Table 2 caption ───────────────────────────────────────────────────
    idx = find_paragraph(doc, "Table 2 Stochastic avoidance failure")
    replace_paragraph_text(doc.paragraphs[idx], (
        "Table 2 Stochastic avoidance failure: outbreak probability by region. "
        "V1 (additive model) and V2 (meth \u00d7 housing interaction model). "
        "Based on n = 2,000 Monte Carlo simulations per model version. V2 "
        "results available for National Average and Pacific Northwest."
    ))

    # ── Additional file 2 description ─────────────────────────────────────
    idx = find_paragraph(doc, "Additional file 2: Stochastic Avoidance")
    replace_paragraph_text(doc.paragraphs[idx], (
        "Additional file 2: Stochastic Avoidance and Outbreak Probability "
        "Analysis. Contains stochastic avoidance conceptual framework, "
        "network density model specification, outbreak probability calculation "
        "methods, model parameters (Table S2.1), national outbreak forecast, "
        "regional heterogeneity analysis, tornado sensitivity analysis, policy "
        "scenario comparisons, V1 versus V2 model comparison with "
        "multiplicative meth \u00d7 housing interaction (Hood et al. 2018), and "
        "supplementary figures (Figs. S1\u2013S11)."
    ))

    # ── Add Reference [21] (Hood et al.) ──────────────────────────────────
    # Find reference 20 (the last existing reference)
    idx = find_paragraph(doc, "20. Grant RM, Lama JR")
    hood_ref = insert_paragraph_after(doc, doc.paragraphs[idx], (
        "21. Hood JE, Buskin SE, Dombrowski JC, Kern DA, Barash EA, "
        "Katzi DA, et al. Methamphetamine use, sexual risk, and decreased "
        "viral suppression among HIV-positive MSM who inject drugs in King "
        "County, Washington. AIDS Patient Care STDs. 2018;32(6):223\u2013233. "
        "https://doi.org/10.1089/apc.2018.0005"
    ))

    # ── Save ──────────────────────────────────────────────────────────────
    doc.save(str(MANUSCRIPT_OUT))
    print(f"Manuscript saved to: {MANUSCRIPT_OUT}")


# ─── COVER LETTER EDITS ──────────────────────────────────────────────────────

def revise_cover_letter():
    doc = Document(str(COVER_IN))

    # ── P10: Key paragraph ────────────────────────────────────────────────
    idx = find_paragraph(doc, "This study addresses a critical gap")
    replace_paragraph_text(doc.paragraphs[idx], (
        "This study addresses a critical gap in HIV prevention research: the "
        "systematic exclusion of people who inject drugs (PWID) from the "
        "benefits of highly efficacious biomedical prevention tools. Using "
        "Monte Carlo simulation of an 8-step LAI-PrEP cascade under current "
        "U.S. policy conditions, we demonstrate a 5,434-fold disparity between "
        "PWID and MSM prevention outcomes\u2014driven almost entirely by structural "
        "barriers (93.1%) rather than pharmacological limitations. We present "
        "two stochastic avoidance models: a baseline framework (V1) predicting "
        "73.0% probability of major outbreak within 5 years, and an enhanced "
        "model (V2) incorporating empirically-derived meth \u00d7 housing interaction "
        "from Hood et al. (2018). Together, V1 and V2 bracket national 5-year "
        "risk at 68\u201373%, while V2 reveals that structural intersections "
        "accelerate Pacific Northwest outbreak timelines to a median of "
        "1.0 year\u2014demonstrating predictable regional system failure rather "
        "than epidemic randomness."
    ))

    # ── P12: Update word count and figure count ───────────────────────────
    idx = find_paragraph(doc, "The manuscript contains approximately")
    replace_paragraph_text(doc.paragraphs[idx], (
        "The manuscript contains approximately 4,200 words, 2 tables, and "
        "7 figures, with 2 additional files providing detailed methodology "
        "and supplementary analyses. All simulation code and data are publicly "
        "available via GitHub."
    ))

    # ── Save ──────────────────────────────────────────────────────────────
    doc.save(str(COVER_OUT))
    print(f"Cover letter saved to: {COVER_OUT}")


# ─── MAIN ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Revising manuscript...")
    revise_manuscript()
    print()
    print("Revising cover letter...")
    revise_cover_letter()
    print()
    print("Done. Revised files:")
    print(f"  {MANUSCRIPT_OUT}")
    print(f"  {COVER_OUT}")
