#!/usr/bin/env python3
"""
Build BMC Additional Files 1 & 2.

Reads data from parameters.json, CSVs, and embeds TIFF figures to produce
two Word documents matching BMC Public Health supplement requirements.

Output:
  /home/claude/HIV_Prevention_PWID/manuscript_prep/BMC_Additional_File_1.docx
  /home/claude/HIV_Prevention_PWID/manuscript_prep/BMC_Additional_File_2.docx
"""

import csv
import json
import os
from pathlib import Path

import math

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def wilson_ci(k, n, alpha=0.05):
    """Wilson score confidence interval for binomial proportion.

    Provides valid coverage for proportions near 0 or 1, unlike the
    normal approximation which can yield negative lower bounds.
    """
    z = 1.96  # for 95% CI
    p_hat = k / n if n > 0 else 0
    denom = 1 + z ** 2 / n
    center = (p_hat + z ** 2 / (2 * n)) / denom
    spread = z * math.sqrt(p_hat * (1 - p_hat) / n + z ** 2 / (4 * n ** 2)) / denom
    return max(0.0, center - spread), min(1.0, center + spread)

# ── Paths ────────────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent
CONFIG = ROOT / "config" / "parameters.json"
BMC_DATA = ROOT / "BMC Public Health" / "data" / "csv_xlsx"
V1_DATA = ROOT / "BMC Public Health" / "Supplemental Data" / "data"
FIG_V1 = ROOT / "SRC" / "MD" / "Data - Results" / "MD_figures_bmc_public_health"
FIG_V2 = ROOT / "BMC Public Health" / "data" / "figures"
OUT_DIR = Path("/home/claude/HIV_Prevention_PWID/manuscript_prep")

# ── Formatting helpers ───────────────────────────────────────────────────────
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(11)
HEADING_SIZE = Pt(12)
SMALL_SIZE = Pt(10)
TABLE_FONT_SIZE = Pt(9)
FIG_WIDTH = Inches(6.5)


def set_run_font(run, size=FONT_SIZE, bold=False, italic=False, color=None):
    """Apply consistent font formatting to a run."""
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Force Times New Roman for East Asian text too
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)


def add_paragraph(doc, text="", bold=False, italic=False, size=FONT_SIZE,
                  alignment=None, space_after=Pt(6), space_before=Pt(0)):
    """Add a formatted paragraph."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading_text(doc, text, level=1):
    """Add a section heading as bold paragraph (not Word heading style)."""
    size = HEADING_SIZE if level == 1 else FONT_SIZE
    return add_paragraph(doc, text, bold=True, size=size,
                         space_before=Pt(12), space_after=Pt(6))


def add_table_caption(doc, text):
    """Add an italic table caption."""
    return add_paragraph(doc, text, italic=True, size=SMALL_SIZE,
                         space_before=Pt(8), space_after=Pt(4))


def format_table(table):
    """Apply consistent formatting to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Style header row
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, size=TABLE_FONT_SIZE, bold=True)
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): 'D9E2F3',
        })
        shading.append(shading_elm)
    # Style data rows
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=TABLE_FONT_SIZE)


def set_cell_text(cell, text, bold=False, align=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    set_run_font(run, size=TABLE_FONT_SIZE, bold=bold)


def add_figure_caption(doc, text):
    """Add figure caption below an image."""
    return add_paragraph(doc, text, italic=True, size=SMALL_SIZE,
                         space_before=Pt(2), space_after=Pt(12))


# ── Data loading ─────────────────────────────────────────────────────────────

def read_data():
    """Load all data sources."""
    data = {}

    # Parameters
    with open(CONFIG) as f:
        data['params'] = json.load(f)

    # Structural barrier results
    with open(BMC_DATA / "structural_barrier_results.csv") as f:
        data['barrier_raw'] = f.read()

    # V2 results
    with open(BMC_DATA / "stochastic_avoidance_v2_results.csv") as f:
        data['v2_raw'] = f.read()

    # Hood comparison
    with open(BMC_DATA / "hood_parameter_comparison_results.csv") as f:
        data['hood_raw'] = f.read()

    # V1 CSVs
    for name in ['national_forecast_summary', 'regional_comparison',
                  'tornado_analysis', 'scenario_comparison']:
        with open(V1_DATA / f"{name}.csv") as f:
            reader = csv.DictReader(f)
            data[f'v1_{name}'] = list(reader)

    return data


def parse_csv_section(raw_text, section_header):
    """Parse a named section from a multi-section CSV file."""
    lines = raw_text.strip().split('\n')
    rows = []
    capturing = False
    header = None
    for line in lines:
        if section_header in line:
            capturing = True
            header = line
            continue
        if capturing:
            if line.strip() == '':
                break
            rows.append(line)
    return header, rows


# ── Additional File 1 ───────────────────────────────────────────────────────

def build_file_1(data):
    """Build Additional File 1: Cascade Model Results."""
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE

    add_header_1(doc)
    add_toc_1(doc)
    add_s1_1_cascade_spec(doc, data)
    add_s1_2_step_probs(doc, data)
    add_s1_3_mc_methods(doc, data)
    add_s1_4_barriers(doc, data)
    add_s1_5_policy(doc, data)
    add_s1_6_snr(doc, data)
    add_s1_7_sensitivity(doc, data)
    add_s1_8_code(doc)
    add_s1_9_refs(doc)

    out_path = OUT_DIR / "BMC_Additional_File_1.docx"
    doc.save(str(out_path))
    print(f"  Saved: {out_path}")
    return out_path


def add_header_1(doc):
    """Title block for Additional File 1."""
    add_paragraph(doc, "ADDITIONAL FILE 1", bold=True, size=Pt(14),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    add_paragraph(doc, "Cascade Model Results, Barrier Decomposition, and Policy Scenarios",
                  bold=True, size=HEADING_SIZE,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    add_paragraph(doc,
                  "Structural Barriers, Stochastic Avoidance, and Outbreak Risk "
                  "in HIV Prevention for PWID",
                  italic=True, size=FONT_SIZE,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_paragraph(doc, "AC Demidont, DO — Nyx Dynamics LLC",
                  size=SMALL_SIZE,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18))


def add_toc_1(doc):
    """Table of contents for File 1."""
    add_heading_text(doc, "Contents")
    sections = [
        "S1.1  Prevention Cascade Model Specification",
        "S1.2  Cascade Step Probabilities",
        "S1.3  Monte Carlo Simulation Methods",
        "S1.4  Barrier Decomposition",
        "S1.5  Policy Scenario Analysis",
        "S1.6  Signal-to-Noise Ratio Analysis",
        "S1.7  Sensitivity Analysis",
        "S1.8  Code Availability",
        "S1.9  References",
    ]
    for s in sections:
        add_paragraph(doc, f"    {s}", size=SMALL_SIZE, space_after=Pt(2))
    add_paragraph(doc, "")  # spacer


def add_s1_1_cascade_spec(doc, data):
    """S1.1 Prevention Cascade Model Specification."""
    add_heading_text(doc, "S1.1  Prevention Cascade Model Specification")
    add_paragraph(doc,
        "The HIV prevention cascade model represents the sequential series of steps "
        "an individual must successfully complete to achieve effective, sustained HIV "
        "prevention. For people who inject drugs (PWID), this cascade is modeled as "
        "an 8-step process, where completion of each step is required before "
        "progressing to the next. Each step has a base probability of success that is "
        "reduced by structural barriers including criminalization/policy, stigma, "
        "infrastructure gaps, and research exclusion.")
    add_paragraph(doc,
        "The cascade completion probability P(cascade) is computed as the product of "
        "all step probabilities:")
    add_paragraph(doc,
        "    P(cascade) = ∏ᵢ₌₁⁸ pᵢ",
        italic=True, space_before=Pt(4), space_after=Pt(4))
    add_paragraph(doc,
        "where pᵢ is the effective probability at step i after barrier penalties are "
        "applied. For PWID, barriers reduce each step's base probability, yielding "
        "dramatically lower cascade completion rates compared to MSM, who face "
        "substantially fewer structural barriers.")


def add_s1_2_step_probs(doc, data):
    """S1.2 Cascade Step Probabilities — Table S1.1."""
    add_heading_text(doc, "S1.2  Cascade Step Probabilities")
    add_paragraph(doc,
        "Table S1.1 presents the cascade step probabilities for PWID and MSM "
        "populations. PWID effective probabilities reflect the base probability "
        "minus all applicable barrier penalties. MSM values represent empirically "
        "observed completion rates from the literature, reflecting a population "
        "with substantially lower structural barriers.")

    add_table_caption(doc,
        "Table S1.1. Prevention cascade step probabilities by population.")

    params = data['params']
    pwid_steps = params['cascade_steps']['pwid']
    msm_steps = params['cascade_steps']['msm']

    step_names = ['awareness', 'willingness', 'healthcare_access', 'disclosure',
                  'provider_willing', 'hiv_testing_adequate', 'first_injection',
                  'sustained_engagement']
    step_labels = [
        'Awareness of PrEP', 'Willingness to use PrEP',
        'Healthcare access', 'Risk disclosure',
        'Provider willing to prescribe', 'Adequate HIV testing',
        'First injection/dose', 'Sustained engagement'
    ]

    table = doc.add_table(rows=1 + len(step_names) + 1, cols=5)
    table.style = 'Table Grid'

    # Header
    headers = ['Step', 'PWID Base', 'PWID Barriers', 'PWID Effective', 'MSM']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    pwid_product = 1.0
    msm_product = 1.0
    for idx, (step, label) in enumerate(zip(step_names, step_labels)):
        row = table.rows[idx + 1]
        step_data = pwid_steps[step]
        base = step_data['base_probability']
        penalties = sum(v for k, v in step_data.items() if k != 'base_probability')
        effective = base - penalties
        msm_val = msm_steps[step]

        set_cell_text(row.cells[0], label)
        set_cell_text(row.cells[1], f"{base:.2f}", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], f"−{penalties:.2f}", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[3], f"{effective:.2f}", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[4], f"{msm_val:.2f}", align=WD_ALIGN_PARAGRAPH.CENTER)
        pwid_product *= effective
        msm_product *= msm_val

    # Cumulative row
    cum_row = table.rows[-1]
    set_cell_text(cum_row.cells[0], "Cumulative P(cascade)", bold=True)
    set_cell_text(cum_row.cells[1], "", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cum_row.cells[2], "", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cum_row.cells[3], f"{pwid_product:.6f}",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cum_row.cells[4], f"{msm_product:.4f}",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    format_table(table)

    add_paragraph(doc,
        f"The cumulative PWID cascade completion probability is "
        f"{pwid_product:.2e}, compared to {msm_product:.4f} for MSM — a "
        f"{msm_product / pwid_product:.0f}-fold disparity driven primarily "
        f"by structural barriers applied at each step.",
        space_before=Pt(8))


def add_s1_3_mc_methods(doc, data):
    """S1.3 Monte Carlo Simulation Methods."""
    add_heading_text(doc, "S1.3  Monte Carlo Simulation Methods")
    params = data['params']
    n = params['simulation']['n_individuals']
    n_sims = params['simulation']['n_stochastic_sims']
    seed = params['simulation']['random_seed']

    add_paragraph(doc,
        f"Monte Carlo simulation was used to model individual-level cascade "
        f"traversal. In each replication, n = {n:,} individuals are simulated "
        f"through the 8-step cascade, with each step modeled as a Bernoulli "
        f"trial using the effective probabilities from Table S1.1.")
    add_paragraph(doc,
        f"Across {n_sims:,} stochastic replications (seed = {seed}), we compute "
        f"the proportion of individuals completing the full cascade and the "
        f"proportion achieving R₀(e) = 0 — complete abrogation of transmission "
        f"risk for a given exposure event (see main text for formal definition). "
        f"This is distinct from the epidemiological basic reproduction number R₀. "
        f"Confidence intervals are computed using the Wilson score method, which "
        f"provides valid coverage for proportions near 0 or 1.")
    add_paragraph(doc,
        "Incarceration disruption is modeled as a competing risk: at each cascade "
        "step, individuals face an annual probability of incarceration that resets "
        "their cascade progress. For PWID, the annual incarceration rate is 30%, "
        "compared to 5% for MSM, further widening the cascade completion gap.")


def add_s1_4_barriers(doc, data):
    """S1.4 Barrier Decomposition — Table S1.2."""
    add_heading_text(doc, "S1.4  Barrier Decomposition")
    add_paragraph(doc,
        "To quantify the relative contribution of different barrier types to "
        "cascade failure, we decompose the total prevention gap into three "
        "primary categories: pathogen biology (intrinsic biological limitations "
        "of prevention tools), HIV testing (inadequate testing infrastructure), "
        "and architectural barriers (structural/systemic obstacles). "
        "Architectural barriers are further decomposed into policy/criminalization, "
        "stigma, infrastructure, and research exclusion sub-categories.")

    add_table_caption(doc,
        "Table S1.2. Barrier decomposition of the PWID HIV prevention cascade gap.")

    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'

    set_cell_text(table.rows[0].cells[0], "Barrier Category", bold=True)
    set_cell_text(table.rows[0].cells[1], "Contribution (%)", bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    barriers = [
        ("Pathogen Biology", "0.0"),
        ("HIV Testing", "6.9"),
        ("Architectural (total)", "93.1"),
        ("    Policy/Criminalization", "38.4"),
        ("    Stigma", "20.6"),
        ("    Infrastructure", "21.9"),
        ("    Research Exclusion", "4.1"),
    ]
    for i, (cat, pct) in enumerate(barriers):
        row = table.rows[i + 1]
        bold = "total" in cat
        set_cell_text(row.cells[0], cat, bold=bold)
        set_cell_text(row.cells[1], pct, bold=bold,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    format_table(table)

    add_paragraph(doc,
        "Architectural barriers account for 93.1% of the PWID prevention gap. "
        "Within architectural barriers, policy/criminalization is the single "
        "largest driver (38.4%), followed by infrastructure gaps (21.9%) and "
        "stigma (20.6%). Research exclusion contributes 4.1%. Pathogen biology "
        "contributes 0.0% because current biomedical tools (PrEP) are "
        "pharmacologically effective for PWID — the failure is entirely in "
        "delivery, not in the intervention itself.",
        space_before=Pt(8))


def add_s1_5_policy(doc, data):
    """S1.5 Policy Scenario Analysis — Table S1.3."""
    add_heading_text(doc, "S1.5  Policy Scenario Analysis")
    add_paragraph(doc,
        "We modeled progressive barrier removal through 8 policy scenarios, "
        "ranging from current policy to theoretical maximum (all barriers "
        "removed). Each scenario adjusts the barrier penalties applied to "
        "cascade steps, and the resulting P(R₀(e) = 0) is computed via Monte "
        "Carlo simulation.")

    add_table_caption(doc,
        "Table S1.3. Policy scenario analysis: cascade completion and P(R₀(e) = 0). "
        "Confidence intervals computed using the Wilson score method.")

    # Parse from structural_barrier_results.csv
    # CSV columns: Scenario[0], Achieved R0=0[1], Completed Cascade[2],
    #   Incarceration Disrupted[3], Observed R0=0 Rate[4], Observed Cascade Rate[5],
    #   95% CI Lower[6], 95% CI Upper[7]
    lines = data['barrier_raw'].strip().split('\n')
    scenario_rows = []
    for line in lines:
        if line.startswith('Scenario,') or line.startswith('Barrier') or \
           line.startswith('Architectural') or line.strip() == '' or \
           line.startswith('Pathogen') or line.startswith('Hiv') or \
           line.startswith('Policy') or line.startswith('Stigma') or \
           line.startswith('Infrastructure') or line.startswith('Research'):
            continue
        parts = line.split(',')
        if len(parts) >= 7:
            # Use parts[4] = Observed R0=0 Rate for all rows (including MSM)
            if parts[4].strip() != '':
                scenario_rows.append(parts)

    table = doc.add_table(rows=1 + len(scenario_rows), cols=4)
    table.style = 'Table Grid'

    headers = ['Scenario', 'P(R₀(e) = 0)', '95% CI', 'Improvement']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    n_individuals = data['params']['simulation']['n_individuals']
    baseline_p = None
    for idx, parts in enumerate(scenario_rows):
        row = table.rows[idx + 1]
        name = parts[0].strip()
        p_r0 = float(parts[4].strip())
        raw_ci_lo = parts[6].strip() if len(parts) > 6 else 'N/A'
        raw_ci_hi = parts[7].strip() if len(parts) > 7 else 'N/A'

        if baseline_p is None:
            baseline_p = p_r0

        if raw_ci_lo == 'N/A':
            ci_str = "N/A"
            improvement = "—"
        else:
            # Recompute CIs using Wilson score method (valid for rare events)
            k = round(p_r0 * n_individuals)
            w_lo, w_hi = wilson_ci(k, n_individuals)
            ci_str = f"[{w_lo:.6f}, {w_hi:.6f}]" if p_r0 < 0.01 else f"[{w_lo:.4f}, {w_hi:.4f}]"
            if baseline_p > 0 and baseline_p != p_r0:
                improvement = f"{p_r0 / baseline_p:.0f}×"
            elif baseline_p == p_r0:
                improvement = "Baseline"
            else:
                improvement = "—"

        set_cell_text(row.cells[0], name)
        set_cell_text(row.cells[1], f"{p_r0:.6f}" if p_r0 < 0.01 else f"{p_r0:.4f}",
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], ci_str, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[3], improvement, align=WD_ALIGN_PARAGRAPH.CENTER)

    format_table(table)

    add_paragraph(doc,
        "Under current policy, P(R₀(e) = 0) is effectively zero (0.003%). "
        "Progressive barrier removal yields stepwise improvement: "
        "decriminalization alone achieves a 66-fold increase, while the "
        "theoretical maximum (all barriers removed) reaches P(R₀(e) = 0) = 19.7%, "
        "approaching the MSM reference of 16.3%. The MSM comparison row "
        "shows that even without any intervention, MSM achieve substantially "
        "higher cascade completion due to lower structural barriers.",
        space_before=Pt(8))


def add_s1_6_snr(doc, data):
    """S1.6 Signal-to-Noise Ratio Analysis — Table S1.4."""
    add_heading_text(doc, "S1.6  Signal-to-Noise Ratio Analysis")
    add_paragraph(doc,
        "The signal-to-noise ratio (SNR) quantifies research efficiency by "
        "comparing the effect size detectable in clinical trials to the "
        "background variability introduced by structural barriers. A higher "
        "SNR indicates that clinical trials can detect treatment effects "
        "more efficiently.")

    add_table_caption(doc,
        "Table S1.4. Signal-to-noise ratio comparison between PWID and MSM populations.")

    snr = data['params']['snr']
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    set_cell_text(table.rows[0].cells[0], "Metric", bold=True)
    set_cell_text(table.rows[0].cells[1], "Value", bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    snr_rows = [
        ("MSM SNR", f"{snr['msm']:,}"),
        ("PWID SNR", f"{snr['pwid']}"),
        ("Disparity", f"{snr['disparity_fold']}-fold"),
    ]
    for i, (label, val) in enumerate(snr_rows):
        set_cell_text(table.rows[i + 1].cells[0], label)
        set_cell_text(table.rows[i + 1].cells[1], val,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    format_table(table)

    add_paragraph(doc,
        f"The MSM research SNR ({snr['msm']:,}) exceeds the PWID SNR "
        f"({snr['pwid']}) by {snr['disparity_fold']}-fold. This disparity "
        f"means that PWID-focused trials require dramatically larger sample "
        f"sizes and longer follow-up periods to detect equivalent treatment "
        f"effects, creating a self-reinforcing cycle of research exclusion.",
        space_before=Pt(8))


def add_s1_7_sensitivity(doc, data):
    """S1.7 Sensitivity Analysis."""
    add_heading_text(doc, "S1.7  Sensitivity Analysis")
    add_paragraph(doc,
        "Probabilistic sensitivity analysis (PSA) was conducted using 1,000 "
        "Latin hypercube samples, varying all barrier penalties simultaneously "
        "within ±50% of their base values. Results confirm the robustness of "
        "the cascade model: across all PSA samples, the PWID cascade completion "
        "rate remained below 0.01%, and the qualitative finding of >90% "
        "architectural barrier contribution was invariant to parameter "
        "perturbation.")
    add_paragraph(doc,
        "One-way sensitivity analysis identified awareness (step 1) as the "
        "most influential cascade step, followed by disclosure and provider "
        "willingness. Fixing any single step to its MSM-equivalent value "
        "improved cascade completion but was insufficient to close the gap, "
        "confirming that the prevention deficit is distributed across multiple "
        "structural barriers rather than concentrated at a single bottleneck.")


def add_s1_8_code(doc):
    """S1.8 Code Availability."""
    add_heading_text(doc, "S1.8  Code Availability")
    add_paragraph(doc,
        "All analysis code, configuration files, and simulation scripts are "
        "available at: https://github.com/Nyx-Dynamics/HIV_Prevention_PWID")
    add_paragraph(doc,
        "Key scripts: structural_barrier_model.py (cascade model and barrier "
        "decomposition), cascade_sensitivity_analysis.py (PSA), and "
        "generate_outputs.py (figure generation).")


def add_s1_9_refs(doc):
    """S1.9 References."""
    add_heading_text(doc, "S1.9  References")
    refs = [
        "[1]  Mayer KH, et al. Antiretroviral pre-exposure prophylaxis implementation "
        "in the United States: a work in progress. J Int AIDS Soc. 2015;18(4 Suppl 3):19980.",
        "[2]  Beyrer C, et al. Global epidemiology of HIV infection in men who have "
        "sex with men. Lancet. 2012;380(9839):367–77.",
        "[3]  Degenhardt L, et al. Global prevalence of injecting drug use and HIV "
        "among people who inject drugs: a multistage systematic review. Lancet. 2017;390:1619–26.",
        "[4]  UNAIDS. Health, Rights and Drugs: Harm Reduction, Decriminalization, "
        "and Zero Discrimination for People Who Use Drugs. Geneva: UNAIDS; 2023.",
        "[5]  Allen ST, et al. Syringe services program utilization and HIV prevention. "
        "Curr Opin HIV AIDS. 2019;14(5):407–15.",
    ]
    for ref in refs:
        add_paragraph(doc, ref, size=SMALL_SIZE, space_after=Pt(2))


# ── Additional File 2 ───────────────────────────────────────────────────────

def build_file_2(data):
    """Build Additional File 2: Stochastic Avoidance and Outbreak Probability."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE

    add_header_2(doc)
    add_toc_2(doc)
    add_s2_1_framework(doc)
    add_s2_2_network_model(doc)
    add_s2_3_outbreak_calc(doc)
    add_s2_4_params(doc, data)
    add_s2_5_v1_national(doc, data)
    add_s2_6_v1_regional(doc, data)
    add_s2_7_v1_tornado(doc, data)
    add_s2_8_v1_scenarios(doc, data)
    add_s2_9_v2_model(doc, data)
    add_s2_10_v2_forecast(doc, data)
    add_s2_11_v2_tornado(doc, data)
    add_s2_12_v2_scenarios(doc, data)
    add_s2_13_comparison(doc, data)
    add_s2_14_hood(doc, data)
    add_s2_15_figures(doc)
    add_s2_16_code(doc)
    add_s2_17_refs(doc)

    out_path = OUT_DIR / "BMC_Additional_File_2.docx"
    doc.save(str(out_path))
    print(f"  Saved: {out_path}")
    return out_path


def add_header_2(doc):
    """Title block for Additional File 2."""
    add_paragraph(doc, "ADDITIONAL FILE 2", bold=True, size=Pt(14),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    add_paragraph(doc,
        "Stochastic Avoidance and Outbreak Probability Analysis",
        bold=True, size=HEADING_SIZE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    add_paragraph(doc,
        "Structural Barriers, Stochastic Avoidance, and Outbreak Risk "
        "in HIV Prevention for PWID",
        italic=True, size=FONT_SIZE,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_paragraph(doc, "AC Demidont, DO — Nyx Dynamics LLC",
                  size=SMALL_SIZE,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18))


def add_toc_2(doc):
    """Table of contents for File 2."""
    add_heading_text(doc, "Contents")
    sections = [
        "S2.1   Stochastic Avoidance Conceptual Framework",
        "S2.2   Network Density Model",
        "S2.3   Outbreak Probability Calculation",
        "S2.4   Model Parameters",
        "S2.5   V1 National Outbreak Forecast",
        "S2.6   V1 Regional Heterogeneity",
        "S2.7   V1 Tornado Analysis",
        "S2.8   V1 Policy Scenarios",
        "S2.9   V2 Model: Meth × Housing Interaction (NEW)",
        "S2.10  V2 National and Regional Forecast (NEW)",
        "S2.11  V2 Tornado Analysis (NEW)",
        "S2.12  V2 Policy Scenarios (NEW)",
        "S2.13  V1 vs V2 Comparison Summary (NEW)",
        "S2.14  Hood et al. Parameter Calibration (NEW)",
        "S2.15  Supplementary Figures",
        "S2.16  Code Availability",
        "S2.17  References",
    ]
    for s in sections:
        add_paragraph(doc, f"    {s}", size=SMALL_SIZE, space_after=Pt(2))
    add_paragraph(doc, "")


def add_s2_1_framework(doc):
    """S2.1 Stochastic Avoidance Conceptual Framework."""
    add_heading_text(doc, "S2.1  Stochastic Avoidance Conceptual Framework")
    add_paragraph(doc,
        "Stochastic avoidance describes the phenomenon by which a disease "
        "outbreak fails to occur despite favorable conditions for sustained "
        "transmission, purely due to random chance in early transmission "
        "events. In the context of HIV among PWID, stochastic avoidance "
        "means that current low outbreak rates may reflect luck rather than "
        "effective prevention.")
    add_paragraph(doc,
        "The key insight is that as network density increases (driven by "
        "methamphetamine co-use, housing instability, and inadequate harm "
        "reduction infrastructure), the probability of stochastic avoidance "
        "decreases and outbreak probability increases. This creates a "
        "time-dependent risk: even if current conditions have not yet "
        "triggered an outbreak, the probability of continued avoidance "
        "declines each year.")


def add_s2_2_network_model(doc):
    """S2.2 Network Density Model."""
    add_heading_text(doc, "S2.2  Network Density Model")
    add_paragraph(doc,
        "We model injection network density as a time-varying function "
        "driven by exogenous factors. The V1 (additive) model computes "
        "effective network density as:")
    add_paragraph(doc,
        "    d(t) = d₀ + α·M(t) + β·H + γ·(1 − SSP) + δ·(1 − OAT)",
        italic=True, space_before=Pt(4), space_after=Pt(4))
    add_paragraph(doc,
        "where d₀ is baseline network density, M(t) is time-varying meth "
        "prevalence, H is housing instability, SSP is syringe service "
        "program coverage, and OAT is opioid agonist therapy coverage. "
        "When d(t) exceeds the critical threshold d*, the local network "
        "is considered capable of sustaining an outbreak.")
    add_paragraph(doc,
        "The V2 model extends this with a multiplicative meth × housing "
        "interaction term (see S2.9).")


def add_s2_3_outbreak_calc(doc):
    """S2.3 Outbreak Probability Calculation."""
    add_heading_text(doc, "S2.3  Outbreak Probability Calculation")
    add_paragraph(doc,
        "Given the time-varying network density d(t), the annual outbreak "
        "probability is modeled as:")
    add_paragraph(doc,
        "    P(outbreak in year t) = p₀ · [d(t) / d*]^κ    if d(t) ≥ d*",
        italic=True, space_before=Pt(4), space_after=Pt(4))
    add_paragraph(doc,
        "where p₀ is the baseline annual outbreak probability, d* is the "
        "critical threshold, and κ is a scaling exponent. The cumulative "
        "outbreak probability over T years is:")
    add_paragraph(doc,
        "    P(outbreak by year T) = 1 − ∏ₜ₌₁ᵀ [1 − P(outbreak in year t)]",
        italic=True, space_before=Pt(4), space_after=Pt(4))
    add_paragraph(doc,
        "Monte Carlo simulation (n = 2,000 per scenario) introduces "
        "stochastic variation in parameters to generate probability "
        "distributions and confidence intervals.")


def add_s2_4_params(doc, data):
    """S2.4 Model Parameters — Table S2.1."""
    add_heading_text(doc, "S2.4  Model Parameters")
    add_paragraph(doc,
        "Table S2.1 presents the V1 stochastic avoidance model parameters. "
        "All values are drawn from published literature or calibrated to "
        "observed outbreak patterns.")

    add_table_caption(doc,
        "Table S2.1. V1 stochastic avoidance model parameters.")

    sa = data['params']['stochastic_avoidance']
    lit = data['params']['literature_params']

    param_rows = [
        ("Baseline network density (d₀)", f"{sa['baseline_network_density']}", "0.08–0.25"),
        ("Meth network multiplier (α)", f"{sa['meth_network_multiplier']}", "1.5–4.0"),
        ("Critical threshold (d*)", f"{sa['critical_threshold']}", "0.25–0.45"),
        ("Baseline outbreak probability (p₀)", f"{sa['baseline_outbreak_prob']}", "0.01–0.08"),
        ("Housing instability (H)", f"{sa['housing_instability']}", "0.55–0.80"),
        ("SSP coverage", f"{sa['ssp_coverage']}", "0.15–0.30"),
        ("OAT coverage", f"{sa['oat_coverage']}", "0.04–0.15"),
        ("Meth prevalence (baseline)", f"{sa['meth_prevalence_baseline']}", "—"),
        ("Meth growth rate (annual)", f"{sa['meth_growth_rate']}", "0.01–0.05"),
        ("Incarceration rate (annual)", f"{data['params']['incarceration']['annual_rate_pwid']}",
         "0.20–0.45"),
    ]

    table = doc.add_table(rows=1 + len(param_rows), cols=3)
    table.style = 'Table Grid'

    headers = ['Parameter', 'Point Estimate', 'Sensitivity Range']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, (name, val, rng) in enumerate(param_rows):
        row = table.rows[idx + 1]
        set_cell_text(row.cells[0], name)
        set_cell_text(row.cells[1], val, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], rng, align=WD_ALIGN_PARAGRAPH.CENTER)

    format_table(table)


def add_s2_5_v1_national(doc, data):
    """S2.5 V1 National Outbreak Forecast — Table S2.2."""
    add_heading_text(doc, "S2.5  V1 National Outbreak Forecast")

    v1 = data['v1_national_forecast_summary']
    # Parse values from standalone CSV
    vals = {row['Metric']: row['Value'] for row in v1}
    n_sims = int(float(vals['n_simulations']))
    median = float(vals['median_years_to_outbreak'])

    # Use comparison-run V1 values as canonical (V1 and V2 run with identical
    # seeds for direct comparison; standalone V1 used a different seed).
    # Comparison CSV: V1 5yr = 0.729938, V1 10yr = 0.933642
    v2_raw = data['v2_raw']
    v1_comp_vals = {}
    in_nat_comp = False
    for line in v2_raw.strip().split('\n'):
        if 'V1 vs V2 NATIONAL COMPARISON' in line:
            in_nat_comp = True
            continue
        if in_nat_comp and line.startswith('Metric,'):
            continue
        if in_nat_comp and ',' in line:
            parts = line.split(',')
            if len(parts) >= 2:
                v1_comp_vals[parts[0].strip()] = parts[1].strip()
        if in_nat_comp and (line.strip() == '' or 'PNW' in line):
            in_nat_comp = False
    p5 = float(v1_comp_vals.get('P(outbreak within 5yr)', vals['p_outbreak_5yr'])) * 100
    p10 = float(v1_comp_vals.get('P(outbreak within 10yr)', vals['p_outbreak_10yr'])) * 100

    add_paragraph(doc,
        f"Under current policy conditions, the V1 model projects substantial "
        f"outbreak risk nationally. Table S2.2 summarizes results from "
        f"n = {n_sims:,} Monte Carlo simulations.")

    add_table_caption(doc,
        "Table S2.2. V1 national outbreak forecast summary.")

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    set_cell_text(table.rows[0].cells[0], "Metric", bold=True)
    set_cell_text(table.rows[0].cells[1], "Value", bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    forecast_rows = [
        ("Number of simulations", f"{n_sims:,}"),
        ("P(outbreak within 5 years)", f"{p5:.1f}%"),
        ("P(outbreak within 10 years)", f"{p10:.1f}%"),
        ("Median years to outbreak", f"{median:.1f}"),
    ]
    for i, (label, val) in enumerate(forecast_rows):
        set_cell_text(table.rows[i + 1].cells[0], label)
        set_cell_text(table.rows[i + 1].cells[1], val,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    format_table(table)

    add_paragraph(doc,
        f"The 5-year outbreak probability of {p5:.1f}% and 10-year probability "
        f"of {p10:.1f}% indicate that continued stochastic avoidance is "
        f"increasingly unlikely. The median time to outbreak of {median:.1f} years "
        f"suggests that, absent intervention, the most likely scenario is an "
        f"outbreak within the near term.",
        space_before=Pt(8))


def add_s2_6_v1_regional(doc, data):
    """S2.6 V1 Regional Heterogeneity — Table S2.3."""
    add_heading_text(doc, "S2.6  V1 Regional Heterogeneity")
    add_paragraph(doc,
        "Outbreak risk varies substantially by region due to differences "
        "in meth prevalence, SSP coverage, housing instability, and "
        "network density.")

    add_table_caption(doc,
        "Table S2.3. V1 regional comparison of outbreak metrics.")

    regions = data['v1_regional_comparison']
    region_labels = {
        'appalachia': 'Appalachia',
        'pacific_northwest': 'Pacific Northwest',
        'northeast_urban': 'Northeast Urban',
        'national_average': 'National Average',
    }

    table = doc.add_table(rows=1 + len(regions), cols=5)
    table.style = 'Table Grid'

    headers = ['Region', 'Outbreak Rate', 'Median (yr)', 'P(5yr)', 'P(10yr)']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, row_data in enumerate(regions):
        row = table.rows[idx + 1]
        region = row_data['Region']
        label = region_labels.get(region, region)
        set_cell_text(row.cells[0], label)
        set_cell_text(row.cells[1],
                      f"{float(row_data['outbreak_rate']):.3f}",
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2],
                      f"{float(row_data['median_years_to_outbreak']):.1f}",
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[3],
                      f"{float(row_data['p_outbreak_5yr']) * 100:.1f}%",
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[4],
                      f"{float(row_data['p_outbreak_10yr']) * 100:.1f}%",
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    format_table(table)

    add_paragraph(doc,
        "The Pacific Northwest shows the highest risk (86.3% at 5 years) "
        "driven by high meth prevalence and network density, while the "
        "national average represents a weighted composite.",
        space_before=Pt(8))


def add_s2_7_v1_tornado(doc, data):
    """S2.7 V1 Tornado Analysis — Table S2.4."""
    add_heading_text(doc, "S2.7  V1 Tornado Analysis")
    add_paragraph(doc,
        "Tornado (one-at-a-time) sensitivity analysis ranks model "
        "parameters by their impact on the 5-year outbreak probability "
        "when varied across their plausible ranges.")

    add_table_caption(doc,
        "Table S2.4. V1 tornado sensitivity analysis — 5-year outbreak probability.")

    tornado = data['v1_tornado_analysis']

    table = doc.add_table(rows=1 + len(tornado), cols=5)
    table.style = 'Table Grid'

    headers = ['Rank', 'Parameter', 'Low Value', 'High Value', 'Range (pp)']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, row_data in enumerate(tornado):
        row = table.rows[idx + 1]
        rng_pp = float(row_data['Outcome Range']) * 100
        set_cell_text(row.cells[0], str(idx + 1),
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], row_data['Parameter'])
        set_cell_text(row.cells[2], row_data['Lower Bound'],
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[3], row_data['Upper Bound'],
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[4], f"{rng_pp:.1f}",
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    format_table(table)


def add_s2_8_v1_scenarios(doc, data):
    """S2.8 V1 Policy Scenarios — Table S2.5."""
    add_heading_text(doc, "S2.8  V1 Policy Scenarios")
    add_paragraph(doc,
        "Table S2.5 presents V1 scenario results across policy "
        "interventions, showing SSP/OAT settings and resulting "
        "5-year outbreak probabilities.")

    add_table_caption(doc,
        "Table S2.5. V1 scenario comparison — policy interventions.")

    scenarios = data['v1_scenario_comparison']

    table = doc.add_table(rows=1 + len(scenarios), cols=5)
    table.style = 'Table Grid'

    headers = ['Scenario', 'SSP Coverage', 'OAT Coverage', 'P(5yr)', 'Median (yr)']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, row_data in enumerate(scenarios):
        row = table.rows[idx + 1]
        set_cell_text(row.cells[0], row_data['Scenario'])
        set_cell_text(row.cells[1], row_data['ssp_coverage'],
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], row_data['oat_coverage'],
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[3],
                      f"{float(row_data['p_outbreak_5yr']) * 100:.1f}%",
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[4],
                      f"{float(row_data['median_years_to_outbreak']):.1f}",
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    format_table(table)

    add_paragraph(doc,
        "Even the most aggressive V1 policy scenario (Full Harm Reduction) "
        "only reduces 5-year outbreak probability from 69.5% to 50.8%, "
        "indicating that outbreak risk remains substantial under all "
        "modeled interventions.",
        space_before=Pt(8))


def add_s2_9_v2_model(doc, data):
    """S2.9 V2 Model: Meth × Housing Interaction — Table S2.6 (NEW)."""
    add_heading_text(doc, "S2.9  V2 Model: Meth × Housing Interaction (NEW)")
    add_paragraph(doc,
        "The V2 model extends V1 by introducing a multiplicative interaction "
        "term between methamphetamine co-use and housing instability. "
        "Epidemiological evidence (Hood et al. 2018 [21]) demonstrates that "
        "the joint effect of meth use and unstable housing on HIV risk "
        "exceeds the sum of individual effects by approximately 1.5×, "
        "consistent with a synergistic interaction.")
    add_paragraph(doc,
        "The V2 network density model becomes:")
    add_paragraph(doc,
        "    d(t) = d₀ + α·M(t) + β·H + λ·M(t)·H + γ·(1 − SSP) + δ·(1 − OAT)",
        italic=True, space_before=Pt(4), space_after=Pt(4))
    add_paragraph(doc,
        "where λ is the meth × housing interaction coefficient.")

    add_table_caption(doc,
        "Table S2.6. V2 model parameters — interaction extension.")

    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'

    headers = ['Parameter', 'Point Estimate', 'Lower Bound', 'Upper Bound']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    set_cell_text(table.rows[1].cells[0], "Meth × housing interaction (λ)")
    set_cell_text(table.rows[1].cells[1], "0.8", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.rows[1].cells[2], "0.3", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.rows[1].cells[3], "1.5", align=WD_ALIGN_PARAGRAPH.CENTER)

    format_table(table)

    add_paragraph(doc,
        "The interaction coefficient λ = 0.8 (range 0.3–1.5) is calibrated "
        "to Hood et al. [21], which found that meth use among PWID with "
        "unstable housing produced joint HIV risk 1.5× greater than the "
        "sum of individual risks. This interaction accelerates network "
        "density growth in regions with high meth prevalence and housing "
        "instability, such as the Pacific Northwest.",
        space_before=Pt(8))


def add_s2_10_v2_forecast(doc, data):
    """S2.10 V2 National and Regional Forecast — Table S2.7 (NEW)."""
    add_heading_text(doc, "S2.10  V2 National and Regional Forecast (NEW)")

    # Parse V2 results
    v2_raw = data['v2_raw']
    lines = v2_raw.strip().split('\n')

    # Extract V2 national values
    nat_section = False
    pnw_section = False
    nat_vals = {}
    pnw_vals = {}

    for line in lines:
        if 'V2 NATIONAL FORECAST SUMMARY' in line:
            nat_section = True
            pnw_section = False
            continue
        if 'V2 PACIFIC NORTHWEST FORECAST SUMMARY' in line:
            nat_section = False
            pnw_section = True
            continue
        if 'V1 vs V2' in line:
            nat_section = False
            pnw_section = False
            continue
        if line.startswith('Metric,'):
            continue
        if nat_section and ',' in line:
            parts = line.split(',', 1)
            if len(parts) == 2:
                nat_vals[parts[0].strip()] = parts[1].strip()
        if pnw_section and ',' in line:
            parts = line.split(',', 1)
            if len(parts) == 2:
                pnw_vals[parts[0].strip()] = parts[1].strip()

    add_paragraph(doc,
        "The V2 model produces modestly different forecasts compared to V1, "
        "with the interaction term primarily affecting regional estimates "
        "where meth and housing co-occur.")

    add_table_caption(doc,
        "Table S2.7. V2 national and Pacific Northwest outbreak forecast.")

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['Metric', 'National', 'Pacific Northwest']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    nat_p5 = float(nat_vals.get('p_outbreak_5yr', 0)) * 100
    nat_p10 = float(nat_vals.get('p_outbreak_10yr', 0)) * 100
    nat_med = float(nat_vals.get('median_years_to_outbreak', 0))
    nat_n = int(float(nat_vals.get('n_simulations', 0)))
    pnw_p5 = float(pnw_vals.get('p_outbreak_5yr', 0)) * 100
    pnw_p10 = float(pnw_vals.get('p_outbreak_10yr', 0)) * 100
    pnw_med = float(pnw_vals.get('median_years_to_outbreak', 0))

    forecast_rows = [
        ("n simulations", f"{nat_n:,}", f"{int(float(pnw_vals.get('n_simulations', 0))):,}"),
        ("P(outbreak within 5yr)", f"{nat_p5:.1f}%", f"{pnw_p5:.1f}%"),
        ("P(outbreak within 10yr)", f"{nat_p10:.1f}%", f"{pnw_p10:.1f}%"),
        ("Median years to outbreak", f"{nat_med:.1f}", f"{pnw_med:.1f}"),
    ]
    for i, (label, nat, pnw) in enumerate(forecast_rows):
        set_cell_text(table.rows[i + 1].cells[0], label)
        set_cell_text(table.rows[i + 1].cells[1], nat,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.rows[i + 1].cells[2], pnw,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    format_table(table)

    add_paragraph(doc,
        f"Nationally, V2 projects {nat_p5:.1f}% 5-year and {nat_p10:.1f}% "
        f"10-year outbreak probabilities, with median time to outbreak of "
        f"{nat_med:.1f} years. The Pacific Northwest shows substantially "
        f"elevated risk ({pnw_p5:.1f}% at 5 years, {pnw_p10:.1f}% at 10 years) "
        f"due to the meth × housing interaction amplifying network density "
        f"in a region with high co-prevalence of both factors.",
        space_before=Pt(8))


def add_s2_11_v2_tornado(doc, data):
    """S2.11 V2 Tornado Analysis — Table S2.8 (NEW)."""
    add_heading_text(doc, "S2.11  V2 Tornado Analysis (NEW)")

    # Parse V2 tornado data
    v2_raw = data['v2_raw']
    lines = v2_raw.strip().split('\n')
    tornado_rows = []
    in_tornado = False
    for line in lines:
        if 'V2 TORNADO SENSITIVITY ANALYSIS' in line:
            in_tornado = True
            continue
        if in_tornado and line.startswith('Rank,'):
            continue
        if in_tornado:
            if line.strip() == '' or 'V2 SCENARIO' in line:
                in_tornado = False
                continue
            parts = line.split(',')
            if len(parts) >= 7:
                tornado_rows.append(parts)

    add_paragraph(doc,
        "The V2 tornado analysis includes the meth × housing interaction "
        "coefficient alongside the original 9 parameters.")

    add_table_caption(doc,
        "Table S2.8. V2 tornado sensitivity analysis — 5-year outbreak probability.")

    table = doc.add_table(rows=1 + len(tornado_rows), cols=5)
    table.style = 'Table Grid'

    headers = ['Rank', 'Parameter', 'Low Value', 'High Value', 'Range (pp)']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, parts in enumerate(tornado_rows):
        row = table.rows[idx + 1]
        set_cell_text(row.cells[0], parts[0].strip(),
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], parts[1].strip())
        set_cell_text(row.cells[2], parts[2].strip(),
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[3], parts[3].strip(),
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[4], parts[6].strip(),
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    format_table(table)

    add_paragraph(doc,
        "The meth × housing interaction coefficient ranks 6th with an "
        "11.2pp range, confirming its meaningful contribution to outbreak "
        "risk. Baseline outbreak probability remains the dominant parameter "
        "(52.7pp range), followed by network density and meth effect.",
        space_before=Pt(8))


def add_s2_12_v2_scenarios(doc, data):
    """S2.12 V2 Policy Scenarios — Table S2.9 (NEW)."""
    add_heading_text(doc, "S2.12  V2 Policy Scenarios (NEW)")

    # Parse V2 scenario data
    v2_raw = data['v2_raw']
    lines = v2_raw.strip().split('\n')
    scenario_rows = []
    in_scenario = False
    for line in lines:
        if 'V2 SCENARIO COMPARISON' in line:
            in_scenario = True
            continue
        if in_scenario and line.startswith('Scenario,'):
            continue
        if in_scenario:
            if line.strip() == '' or 'V2 PROBABILISTIC' in line:
                in_scenario = False
                continue
            parts = line.split(',')
            if len(parts) >= 3:
                scenario_rows.append(parts)

    add_paragraph(doc,
        "V2 policy scenarios incorporate the meth × housing interaction, "
        "producing systematically different scenario outcomes compared to V1.")

    add_table_caption(doc,
        "Table S2.9. V2 scenario comparison — policy interventions.")

    table = doc.add_table(rows=1 + len(scenario_rows), cols=3)
    table.style = 'Table Grid'

    headers = ['Scenario', 'P(5yr Outbreak)', 'Median Years']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, parts in enumerate(scenario_rows):
        row = table.rows[idx + 1]
        set_cell_text(row.cells[0], parts[0].strip())
        p5 = float(parts[1].strip()) * 100
        set_cell_text(row.cells[1], f"{p5:.1f}%",
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], parts[2].strip(),
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    format_table(table)

    add_paragraph(doc,
        "Under V2 current policy, the 5-year outbreak probability is 65.0%. "
        "Full harm reduction reduces this to 44.6% — a meaningful reduction "
        "but still unacceptably high, underscoring that even comprehensive "
        "policy interventions cannot eliminate outbreak risk in the presence "
        "of synergistic meth × housing effects.",
        space_before=Pt(8))


def add_s2_13_comparison(doc, data):
    """S2.13 V1 vs V2 Comparison Summary — Table S2.10 (NEW)."""
    add_heading_text(doc, "S2.13  V1 vs V2 Comparison Summary (NEW)")

    # Parse comparison data
    v2_raw = data['v2_raw']
    lines = v2_raw.strip().split('\n')
    nat_comp = []
    pnw_comp = []
    in_nat = False
    in_pnw = False

    for line in lines:
        if 'V1 vs V2 NATIONAL COMPARISON' in line:
            in_nat = True
            in_pnw = False
            continue
        if 'V1 vs V2 PNW COMPARISON' in line:
            in_nat = False
            in_pnw = True
            continue
        if 'V2 TORNADO' in line:
            in_pnw = False
            continue
        if line.startswith('Metric,'):
            continue
        if in_nat and ',' in line:
            parts = line.split(',')
            if len(parts) >= 4:
                nat_comp.append(parts)
        if in_pnw and ',' in line:
            parts = line.split(',')
            if len(parts) >= 4:
                pnw_comp.append(parts)

    add_paragraph(doc,
        "Table S2.10 provides a side-by-side comparison of V1 and V2 "
        "predictions at both national and Pacific Northwest levels.")

    add_table_caption(doc,
        "Table S2.10. V1 vs V2 comparison — national and Pacific Northwest.")

    # Build combined table: 1 header + 1 nat sub-header + nat rows + 1 pnw sub-header + pnw rows
    table = doc.add_table(rows=1 + 1 + len(nat_comp) + 1 + len(pnw_comp), cols=4)
    table.style = 'Table Grid'

    headers = ['Metric', 'V1', 'V2', 'Δ']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    row_idx = 1
    # National section header
    nat_header_row = table.rows[row_idx]
    set_cell_text(nat_header_row.cells[0], "National", bold=True)
    for c in range(1, 4):
        set_cell_text(nat_header_row.cells[c], "")
    # Merge or shade
    for cell in nat_header_row.cells:
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'E2EFDA'
        })
        shading.append(shading_elm)
    row_idx += 1

    for parts in nat_comp:
        row = table.rows[row_idx]
        metric = parts[0].strip()
        v1_val = parts[1].strip()
        v2_val = parts[2].strip()
        delta = parts[3].strip()

        # Format percentages
        if 'P(outbreak' in metric:
            v1_pct = f"{float(v1_val) * 100:.1f}%"
            v2_pct = f"{float(v2_val) * 100:.1f}%"
            d_pct = f"{float(delta) * 100:+.1f}pp"
        elif 'Median' in metric:
            v1_pct = f"{float(v1_val):.1f}"
            v2_pct = f"{float(v2_val):.1f}"
            d_pct = f"{float(delta):+.1f}"
        elif 'rate' in metric.lower():
            v1_pct = f"{float(v1_val) * 100:.1f}%"
            v2_pct = f"{float(v2_val) * 100:.1f}%"
            d_pct = f"{float(delta) * 100:+.1f}pp"
        else:
            v1_pct = v1_val
            v2_pct = v2_val
            d_pct = delta

        set_cell_text(row.cells[0], f"  {metric}")
        set_cell_text(row.cells[1], v1_pct, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], v2_pct, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[3], d_pct, align=WD_ALIGN_PARAGRAPH.CENTER)
        row_idx += 1

    # PNW section - we need to add more rows if needed
    # The table was pre-allocated, now fill PNW rows
    # Actually we allocated: 1 header + len(nat_comp) + 1 pnw_header + len(pnw_comp)
    pnw_header_row = table.rows[row_idx]
    set_cell_text(pnw_header_row.cells[0], "Pacific Northwest", bold=True)
    for c in range(1, 4):
        set_cell_text(pnw_header_row.cells[c], "")
    for cell in pnw_header_row.cells:
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'E2EFDA'
        })
        shading.append(shading_elm)
    row_idx += 1

    for parts in pnw_comp:
        row = table.rows[row_idx]
        metric = parts[0].strip()
        v1_val = parts[1].strip()
        v2_val = parts[2].strip()
        delta = parts[3].strip()

        if 'P(outbreak' in metric:
            v1_pct = f"{float(v1_val) * 100:.1f}%"
            v2_pct = f"{float(v2_val) * 100:.1f}%"
            d_pct = f"{float(delta) * 100:+.1f}pp"
        elif 'Median' in metric:
            v1_pct = f"{float(v1_val):.1f}"
            v2_pct = f"{float(v2_val):.1f}"
            d_pct = f"{float(delta):+.1f}"
        elif 'rate' in metric.lower():
            v1_pct = f"{float(v1_val) * 100:.1f}%"
            v2_pct = f"{float(v2_val) * 100:.1f}%"
            d_pct = f"{float(delta) * 100:+.1f}pp"
        else:
            v1_pct = v1_val
            v2_pct = v2_val
            d_pct = delta

        set_cell_text(row.cells[0], f"  {metric}")
        set_cell_text(row.cells[1], v1_pct, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], v2_pct, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[3], d_pct, align=WD_ALIGN_PARAGRAPH.CENTER)
        row_idx += 1

    format_table(table)

    add_paragraph(doc,
        "At the national level, V2 produces slightly lower outbreak "
        "probabilities than V1 (−4.6pp at 5 years) because the interaction "
        "term redistributes risk toward high-prevalence regions. The Pacific "
        "Northwest shows increased risk under V2 (+5.7pp at 5 years), "
        "reflecting the region's high meth × housing co-prevalence.",
        space_before=Pt(8))


def add_s2_14_hood(doc, data):
    """S2.14 Hood et al. Parameter Calibration — Table S2.11 (NEW)."""
    add_heading_text(doc, "S2.14  Hood et al. Parameter Calibration (NEW)")
    add_paragraph(doc,
        "Hood et al. (2018) [21] provided empirical data from the King County "
        "(Seattle) PWID cohort that was used to calibrate model parameters. "
        "Table S2.11 presents a 3-way comparison of parameter values across "
        "the original model, Hood-adjusted V1, and V2+Hood configurations.")

    add_table_caption(doc,
        "Table S2.11. Hood et al. parameter calibration — 3-way comparison.")

    # Parse hood data
    hood_raw = data['hood_raw']
    lines = hood_raw.strip().split('\n')
    param_rows = []
    in_params = False
    for line in lines:
        if 'HOOD PARAMETER ADJUSTMENTS' in line:
            in_params = True
            continue
        if line.startswith('Parameter,') and in_params:
            continue
        if in_params:
            if line.strip() == '' or 'V2 NEW' in line:
                in_params = False
                continue
            parts = line.split(',')
            if len(parts) >= 5:
                param_rows.append(parts)

    table = doc.add_table(rows=1 + len(param_rows), cols=5)
    table.style = 'Table Grid'

    headers = ['Parameter', 'Original', 'Hood-Adjusted', '% Change', 'Source']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    for idx, parts in enumerate(param_rows):
        row = table.rows[idx + 1]
        set_cell_text(row.cells[0], parts[0].strip())
        set_cell_text(row.cells[1], parts[1].strip(),
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], parts[2].strip(),
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[3], parts[3].strip(),
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[4], parts[4].strip())

    format_table(table)

    # Add 3-way forecast comparison text
    add_paragraph(doc,
        "The Hood-calibrated parameters produce broadly consistent results "
        "with the original model at the national level, with V1 Hood showing "
        "+1.3pp and V2+Hood showing −4.3pp change in 5-year outbreak "
        "probability. At the PNW level, V2+Hood shows increased risk "
        "(+2.8pp) due to the interaction term capturing the empirically "
        "observed meth × housing synergy in Seattle-area PWID populations.",
        space_before=Pt(8))


def add_s2_15_figures(doc):
    """S2.15 Supplementary Figures — embed all 11 TIFFs."""
    add_heading_text(doc, "S2.15  Supplementary Figures")
    add_paragraph(doc,
        "Figures S1–S4 present V1 results; Figures S5–S8 present V2 results; "
        "Figures S9–S11 present 3-way (Original/Hood V1/V2+Hood) comparisons.")

    figures = [
        (FIG_V1 / "FigS1_ContextualStochasticFailureDriver.tif",
         "Figure S1. Contextual stochastic failure driver. Projected "
         "trajectories of the exogenous network-density modifier under "
         "current policy conditions."),
        (FIG_V1 / "FigS2_OutbreakForecast.tif",
         "Figure S2. V1 national outbreak forecast. (a) Cumulative "
         "probability reaching 73.0% at 5 years and 93.4% at 10 years. "
         "(b) Time-to-outbreak distribution with median 3.0 years."),
        (FIG_V1 / "FigS3_TornadoDiagram.tif",
         "Figure S3. V1 tornado diagram. Sensitivity of 5-year outbreak "
         "probability to parameter variation across plausible ranges."),
        (FIG_V1 / "FigS4_ScenarioComparison.tif",
         "Figure S4. V1 scenario comparison. 5-year outbreak probability "
         "across policy interventions from current policy to full harm "
         "reduction."),
        (FIG_V2 / "FigS5_V2_NetworkDensityComparison.tif",
         "Figure S5. V1 vs V2 network density comparison. Trajectories "
         "under additive (V1) vs multiplicative interaction (V2) models."),
        (FIG_V2 / "FigS6_V2_OutbreakForecastComparison.tif",
         "Figure S6. V1 vs V2 outbreak forecast comparison. Cumulative "
         "probability curves with simulation intervals."),
        (FIG_V2 / "FigS7_V2_TornadoDiagram.tif",
         "Figure S7. V2 tornado diagram. Meth × housing interaction "
         "ranks 6th with 11.2pp range."),
        (FIG_V2 / "FigS8_V2_ScenarioComparison.tif",
         "Figure S8. V2 scenario comparison. Current Policy 65.0% → "
         "Full Harm Reduction 44.6%."),
        (FIG_V2 / "FigS9_3Way_NetworkDensity.tif",
         "Figure S9. 3-way network density comparison (Original, Hood V1, "
         "V2+Hood)."),
        (FIG_V2 / "FigS10_3Way_OutbreakForecast.tif",
         "Figure S10. 3-way outbreak forecast comparison."),
        (FIG_V2 / "FigS11_3Way_ScenarioComparison.tif",
         "Figure S11. 3-way scenario comparison."),
    ]

    for fig_path, caption in figures:
        if fig_path.exists():
            doc.add_picture(str(fig_path), width=FIG_WIDTH)
            # Center the image
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_figure_caption(doc, caption)
        else:
            add_paragraph(doc,
                f"[Figure not found: {fig_path.name}]",
                italic=True, size=SMALL_SIZE)
            add_figure_caption(doc, caption)


def add_s2_16_code(doc):
    """S2.16 Code Availability."""
    add_heading_text(doc, "S2.16  Code Availability")
    add_paragraph(doc,
        "All analysis code, configuration files, and simulation scripts are "
        "available at: https://github.com/Nyx-Dynamics/HIV_Prevention_PWID")
    add_paragraph(doc,
        "Key scripts: stochastic_avoidance_enhanced.py (V1 model), "
        "stochastic_avoidance_v2.py (V2 model with meth × housing "
        "interaction), hood_parameter_comparison.py (Hood et al. calibration), "
        "and generate_outputs.py (figure generation).")


def add_s2_17_refs(doc):
    """S2.17 References."""
    add_heading_text(doc, "S2.17  References")
    refs = [
        "[3]  Degenhardt L, et al. Global prevalence of injecting drug use and HIV "
        "among people who inject drugs. Lancet. 2017;390:1619–26.",
        "[4]  UNAIDS. Health, Rights and Drugs. Geneva: UNAIDS; 2023.",
        "[9]  Van Handel MM, et al. County-level vulnerability assessment for rapid "
        "dissemination of HIV or HCV infections among persons who inject drugs, "
        "United States. J Acquir Immune Defic Syndr. 2016;73(3):323–31.",
        "[10] Peters PJ, et al. HIV infection linked to injection use of oxymorphone "
        "in Indiana, 2014–2015. N Engl J Med. 2016;375(3):229–39.",
        "[14] Cranston K, et al. Notes from the field: HIV diagnoses among persons "
        "who inject drugs — Northeastern Massachusetts, 2015–2018. MMWR. "
        "2019;68(10):253–4.",
        "[15] Ondocsin J, et al. An outbreak of HIV among people who inject drugs "
        "in West Virginia. N Engl J Med. 2023;388(15):1377–83.",
        "[16] Golden MR, et al. Outbreak of human immunodeficiency virus infection "
        "among heterosexual persons who are living homeless and who smoke "
        "fentanyl — King County, Washington, 2023. MMWR. 2024;73(6):141–6.",
        "[17] Des Jarlais DC, et al. HIV among people who inject drugs: a 25-year "
        "perspective on the Amsterdam, New York City, and 'prevent AIDS' studies. "
        "Drug Alcohol Depend. 2021;225:108788.",
        "[18] Strathdee SA, et al. HIV and risk environment for injecting drug users: "
        "the past, present, and future. Lancet. 2010;376(9737):268–84.",
        "[21] Hood JE, et al. Engagement in HIV care among persons who inject drugs "
        "in King County, Washington: associations with methamphetamine use and "
        "housing status. AIDS Patient Care STDS. 2018;32(4):127–34.",
    ]
    for ref in refs:
        add_paragraph(doc, ref, size=SMALL_SIZE, space_after=Pt(2))


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    """Build both additional files."""
    print("Building BMC Additional Files...")
    print()

    # Ensure output directory exists
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    # Load data
    print("Loading data sources...")
    data = read_data()
    print("  Data loaded successfully.")
    print()

    # Build File 1
    print("Building Additional File 1...")
    build_file_1(data)
    print()

    # Build File 2
    print("Building Additional File 2...")
    build_file_2(data)
    print()

    print("Done! Both files saved to:", OUT_DIR)


if __name__ == "__main__":
    main()
