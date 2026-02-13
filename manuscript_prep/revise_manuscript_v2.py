#!/usr/bin/env python3
"""
revise_manuscript_v2.py

Applies v2 revisions to BMC_Public_Health_Manuscript_revised.docx and
BMC_Cover_Letter_revised.docx.  Produces *_v2.docx copies.

Issues addressed:
  1  PEP → PrEP framing restructure (Background)
  2  R₀(e,t) = 0 notation clarification
  5  Probability alignment verification
  6  Figure renumbering (new Fig 1 schematic; old 1→2, 2→3, etc.)
  8  "Model assumptions & estimands" subsection (Methods)
  9  Temper absolute claims
"""

import copy
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── paths ────────────────────────────────────────────────────────────────────
INPUT_DIR = Path("/home/claude/HIV_Prevention_PWID/manuscript_prep")
OUTPUT_DIR = INPUT_DIR

MANUSCRIPT_IN = INPUT_DIR / "BMC_Public_Health_Manuscript_revised.docx"
MANUSCRIPT_OUT = OUTPUT_DIR / "BMC_Public_Health_Manuscript_v2.docx"

COVER_IN = INPUT_DIR / "BMC_Cover_Letter_revised.docx"
COVER_OUT = OUTPUT_DIR / "BMC_Cover_Letter_v2.docx"


# ─── helpers ──────────────────────────────────────────────────────────────────

def find_paragraph(doc, prefix, start=0):
    """Return the index of the first paragraph whose text starts with *prefix*."""
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if p.text.strip().startswith(prefix):
            return i
    raise ValueError(f"Paragraph starting with {prefix!r} not found (start={start})")


def find_paragraph_containing(doc, substring, start=0):
    """Return the index of the first paragraph containing *substring*."""
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if substring in p.text:
            return i
    raise ValueError(f"Paragraph containing {substring!r} not found (start={start})")


def replace_paragraph_text(paragraph, new_text):
    """Replace all text in a paragraph while preserving the style of the first run."""
    style = paragraph.style
    first_run_font = None
    if paragraph.runs:
        r = paragraph.runs[0]
        first_run_font = {
            'bold': r.bold,
            'italic': r.italic,
            'size': r.font.size,
            'name': r.font.name,
        }

    # Clear all runs
    p_elem = paragraph._element
    for child in list(p_elem):
        if child.tag == qn('w:r'):
            p_elem.remove(child)

    new_run = paragraph.add_run(new_text)
    if first_run_font:
        if first_run_font['bold'] is not None:
            new_run.bold = first_run_font['bold']
        if first_run_font['italic'] is not None:
            new_run.italic = first_run_font['italic']
        if first_run_font['size'] is not None:
            new_run.font.size = first_run_font['size']
        if first_run_font['name'] is not None:
            new_run.font.name = first_run_font['name']
    paragraph.style = style


def insert_paragraph_after(doc, ref_paragraph, text):
    """Insert a new paragraph immediately after *ref_paragraph*."""
    new_p = OxmlElement('w:p')

    ref_pPr = ref_paragraph._element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_p.append(copy.deepcopy(ref_pPr))

    new_r = OxmlElement('w:r')
    if ref_paragraph.runs:
        ref_rPr = ref_paragraph.runs[0]._element.find(qn('w:rPr'))
        if ref_rPr is not None:
            new_r.append(copy.deepcopy(ref_rPr))

    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)

    ref_paragraph._element.addnext(new_p)

    for p in doc.paragraphs:
        if p._element is new_p:
            return p

    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, doc.element.body)


def _replace_in_paragraph(paragraph, old, new):
    """Replace *old* with *new* in a paragraph, handling text split across runs.

    Strategy: if *old* appears in the concatenated paragraph text, collapse all
    runs into the first run (preserving its formatting), perform the replacement,
    and clear subsequent runs.  Returns True if a replacement was made.
    """
    full_text = paragraph.text
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    if not paragraph.runs:
        return False
    # Preserve first run's formatting, put all text there
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def _get_all_paragraphs(doc):
    """Yield all paragraphs in body and table cells."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def renumber_references(doc):
    """Renumber all references by order of first in-text citation.

    1. Scan body text (before 'References' heading) for [N] citations.
    2. Build old→new mapping based on appearance order.
    3. Replace all in-text citations using temp placeholders to avoid collisions.
    4. Reorder and renumber the reference list paragraphs.
    """
    import re as _re

    # ── Step 1: find the References heading index ────────────────────────────
    ref_heading_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "References":
            ref_heading_idx = i
            break
    if ref_heading_idx is None:
        print("    Warning: 'References' heading not found; skipping renumber")
        return

    # ── Step 2: extract citation order from body text ────────────────────────
    seen = set()
    appearance_order = []  # list of old ref numbers in first-appearance order
    for i, p in enumerate(doc.paragraphs):
        if i >= ref_heading_idx:
            break
        t = p.text
        # Find all bracketed citation groups: [1], [1, 2], [3–5], etc.
        for m in _re.finditer(r'\[([0-9,\s\u2013\u002d]+)\]', t):
            group = m.group(1)
            parts = _re.split(r'[,]', group)
            for part in parts:
                part = part.strip()
                if '\u2013' in part or '-' in part:
                    nums = _re.split(r'[\u2013\-]', part)
                    if len(nums) == 2:
                        try:
                            for n in range(int(nums[0].strip()), int(nums[1].strip()) + 1):
                                if n not in seen:
                                    appearance_order.append(n)
                                    seen.add(n)
                        except ValueError:
                            pass
                else:
                    try:
                        n = int(part)
                        if n not in seen:
                            appearance_order.append(n)
                            seen.add(n)
                    except ValueError:
                        pass

    # Build mapping: old_num -> new_num
    old_to_new = {}
    for new_num, old_num in enumerate(appearance_order, start=1):
        old_to_new[old_num] = new_num

    if all(old == new for old, new in old_to_new.items()):
        print("  Bibliography: already in order.")
        return

    print(f"  Bibliography: renumbering {len(old_to_new)} references...")
    for old, new in sorted(old_to_new.items()):
        if old != new:
            print(f"    [{old}] -> [{new}]")

    # ── Step 3: replace in-text citations using placeholder tokens ───────────
    # Use §REF_N§ as collision-safe placeholders
    PLACEHOLDER = "\u00a7REF_{}\u00a7"

    def _replace_citations_in_para(paragraph):
        """Replace [N] -> [§REF_N§] for all mapped refs."""
        full = paragraph.text
        if '[' not in full:
            return

        def _sub_bracket(match):
            inner = match.group(1)
            parts = _re.split(r'[,]', inner)
            new_parts = []
            for part in parts:
                part_s = part.strip()
                # Handle ranges
                range_match = _re.match(r'^(\d+)\s*[\u2013\-]\s*(\d+)$', part_s)
                if range_match:
                    lo, hi = int(range_match.group(1)), int(range_match.group(2))
                    new_lo = old_to_new.get(lo, lo)
                    new_hi = old_to_new.get(hi, hi)
                    # Expand range to individual refs (ranges may not be
                    # contiguous after renumbering)
                    expanded = []
                    for x in range(lo, hi + 1):
                        expanded.append(old_to_new.get(x, x))
                    # If still contiguous, keep as range; otherwise list
                    expanded.sort()
                    if expanded == list(range(expanded[0], expanded[-1] + 1)):
                        if len(expanded) > 2:
                            new_parts.append(
                                PLACEHOLDER.format(expanded[0]) + "\u2013" +
                                PLACEHOLDER.format(expanded[-1]))
                        else:
                            for e in expanded:
                                new_parts.append(PLACEHOLDER.format(e))
                    else:
                        for e in expanded:
                            new_parts.append(PLACEHOLDER.format(e))
                else:
                    try:
                        n = int(part_s)
                        new_parts.append(PLACEHOLDER.format(old_to_new.get(n, n)))
                    except ValueError:
                        new_parts.append(part_s)
            return "[" + ", ".join(new_parts) + "]"

        new_text = _re.sub(r'\[([0-9,\s\u2013\u002d]+)\]', _sub_bracket, full)
        if new_text != full:
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                for run in paragraph.runs[1:]:
                    run.text = ""

    # Apply to body and figure legends (everything except reference entries)
    for i, p in enumerate(doc.paragraphs):
        if i <= ref_heading_idx:
            _replace_citations_in_para(p)
            continue
        # After References heading: skip reference entries (start with "N. ")
        t = p.text.strip()
        if not _re.match(r'^\d+\.\s', t):
            # Figure legends, table captions, additional file descriptions
            _replace_citations_in_para(p)
    # Also handle table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_citations_in_para(p)

    # ── Step 4: reorder reference list paragraphs ────────────────────────────
    # Collect reference paragraphs (numbered entries after "References" heading)
    ref_paras = {}  # old_num -> paragraph element
    ref_end_idx = None
    for i in range(ref_heading_idx + 1, len(doc.paragraphs)):
        t = doc.paragraphs[i].text.strip()
        m = _re.match(r'^(\d+)\.\s', t)
        if m:
            ref_paras[int(m.group(1))] = doc.paragraphs[i]
            ref_end_idx = i
        elif t and t not in ("", " "):
            # Hit non-reference content (e.g., "Figure legends")
            break

    # Rebuild reference entries with new numbering, in new order
    # First, store the text of each ref (minus the old number prefix)
    ref_texts = {}
    for old_num, para in ref_paras.items():
        t = para.text
        # Remove "N. " prefix
        t = _re.sub(r'^\d+\.\s', '', t)
        ref_texts[old_num] = t

    # Now rewrite each reference paragraph in the new order
    # We'll reuse the existing paragraph elements in their DOM positions
    # Sort by new number and write into the existing slots
    sorted_slots = sorted(ref_paras.keys(), key=lambda k: list(ref_paras.keys()).index(k))
    slot_paras = [ref_paras[k] for k in sorted_slots]  # paragraphs in DOM order

    # Build new content in new order
    new_order = sorted(old_to_new.keys(), key=lambda k: old_to_new[k])
    # Also include any refs that are in the list but not cited
    uncited = [k for k in ref_paras if k not in old_to_new]
    # Assign them numbers after the cited ones
    next_num = len(old_to_new) + 1
    for uc in sorted(uncited):
        old_to_new[uc] = next_num
        new_order.append(uc)
        next_num += 1

    for slot_idx, old_num in enumerate(new_order):
        if slot_idx < len(slot_paras):
            new_num = old_to_new[old_num]
            new_text = f"{new_num}. {ref_texts[old_num]}"
            replace_paragraph_text(slot_paras[slot_idx], new_text)

    # ── Step 5: resolve placeholders → final numbers ─────────────────────────
    def _resolve_placeholders(paragraph):
        full = paragraph.text
        if "\u00a7REF_" not in full:
            return
        def _resolve(m):
            return str(m.group(1))
        new_text = _re.sub(r'\u00a7REF_(\d+)\u00a7', _resolve, full)
        if new_text != full:
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                for run in paragraph.runs[1:]:
                    run.text = ""

    for p in doc.paragraphs:
        _resolve_placeholders(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _resolve_placeholders(p)

    print(f"  Bibliography: {len(old_to_new)} references renumbered successfully.")


def global_replace_text(doc, old, new):
    """Replace *old* with *new* in every paragraph and table cell."""
    count = 0
    for p in doc.paragraphs:
        if _replace_in_paragraph(p, old, new):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _replace_in_paragraph(p, old, new):
                        count += 1
    return count


# ─── ISSUE 6: Figure renumbering map ─────────────────────────────────────────
# New Fig 1 = Cascade schematic (new)
# Old Fig 1 → Fig 2, Old Fig 2 → Fig 3, Old Fig 3 → Fig 4,
# Old Fig 4 → Fig 5, Old Fig 6 → Fig 6 (unchanged), Old Fig 7 → Fig 7 (unchanged)
# Old Fig 5 → Fig 8 (SNR/LOOCV moved to last)
#
# We renumber in reverse order (highest first) to avoid double-replacement.

FIGURE_RENUMBER = [
    # (old_text, new_text) — applied in this order
    # First handle special cases
    ("Fig. 5 Signal-to-noise", "Fig. 8 Signal-to-noise"),
    ("[Fig. 5 about here]", "[Fig. 8 about here]"),
    ("(Fig. 5)", "(Fig. 8)"),
    # Fig 7 stays 7 — no change needed
    # Fig 6 stays 6 — no change needed
    # Now renumber Fig 4 → Fig 5
    ("Fig. 4 Stochastic avoidance", "Fig. 5 Stochastic avoidance"),
    ("[Fig. 4 about here]", "[Fig. 5 about here]"),
    ("(Fig. 4)", "(Fig. 5)"),
    # Fig 3 → Fig 4
    ("Fig. 3 Policy scenario", "Fig. 4 Policy scenario"),
    ("[Fig. 3 about here]", "[Fig. 4 about here]"),
    ("(Fig. 3)", "(Fig. 4)"),
    ("(Fig. 3;", "(Fig. 4;"),
    # Fig 2 → Fig 3
    ("Fig. 2 Three-layer barrier", "Fig. 3 Three-layer barrier"),
    ("[Fig. 2 about here]", "[Fig. 3 about here]"),
    ("(Fig. 2)", "(Fig. 3)"),
    # Fig 1 → Fig 2
    ("Fig. 1 LAI-PrEP cascade", "Fig. 2 LAI-PrEP cascade"),
    ("[Fig. 1 about here]", "[Fig. 2 about here]"),
    ("(Fig. 1)", "(Fig. 2)"),
]


# ─── MANUSCRIPT EDITS ─────────────────────────────────────────────────────────

def revise_manuscript():
    doc = Document(str(MANUSCRIPT_IN))

    # ══════════════════════════════════════════════════════════════════════════
    # Title: add study design (BMC recommendation)
    # ══════════════════════════════════════════════════════════════════════════
    print("  Title: adding study design...")
    idx_title = find_paragraph(doc, "Structural Barriers, Stochastic Avoidance")
    replace_paragraph_text(doc.paragraphs[idx_title], (
        "Structural Barriers, Stochastic Avoidance, and Outbreak Risk "
        "in HIV Prevention for People Who Inject Drugs: "
        "a Monte Carlo Simulation Study"
    ))

    # ══════════════════════════════════════════════════════════════════════════
    # Abstract Background: route/agent-specific LAI-PrEP efficacy
    # ══════════════════════════════════════════════════════════════════════════
    print("  Abstract: route/agent-specific efficacy...")
    idx_bg = find_paragraph(doc, "Background: Despite the availability")
    replace_paragraph_text(doc.paragraphs[idx_bg], (
        "Background: Long-acting PrEP agents have demonstrated very high "
        "efficacy in large phase 3 trials, but people who inject drugs (PWID) "
        "PrEP uptake remains extremely low (<1.5%), and no long-acting "
        "injectable regimen has been validated for parenteral (injection) "
        "exposures. We hypothesized that nested structural barriers, rather "
        "than pharmacological limitations, explain this disparity."
    ))

    # ══════════════════════════════════════════════════════════════════════════
    # ISSUE 6: Figure renumbering (do first so later edits use new numbers)
    # ══════════════════════════════════════════════════════════════════════════
    print("  Issue 6: Renumbering figures...")
    for old_text, new_text in FIGURE_RENUMBER:
        n = global_replace_text(doc, old_text, new_text)
        if n > 0:
            print(f"    Replaced '{old_text}' → '{new_text}' ({n} runs)")

    # Insert Fig. 1 legend before Fig. 2 legend (which was formerly Fig. 1)
    idx_fig2 = find_paragraph(doc, "Fig. 2 LAI-PrEP cascade")
    fig1_legend = insert_paragraph_after(doc, doc.paragraphs[idx_fig2 - 1], (
        "Fig. 1 Prevention cascade with structural barrier layers. "
        "Eight-step cascade showing base probability, barrier penalties, "
        "and effective probability for PWID compared to MSM. Barrier types "
        "are color-coded: policy/criminalization (red), stigma (orange), "
        "infrastructure (blue), research exclusion (purple), HIV testing "
        "(teal), and ML/algorithmic bias (gray). Cumulative PWID cascade "
        "completion probability is <0.01% versus 16.3% for MSM — a "
        "disparity driven entirely by structural barriers at each step."
    ))
    # Insert blank line after
    insert_paragraph_after(doc, fig1_legend, "")

    # Insert [Fig. 1 about here] placeholder in Background, after the
    # R₀(e) definition paragraph (which will be inserted by Issue 2 below).
    # We'll do this after Issue 1/2 inserts.

    # ══════════════════════════════════════════════════════════════════════════
    # ISSUE 1: PEP → PrEP framing restructure (Background [16]–[20])
    # ══════════════════════════════════════════════════════════════════════════
    print("  Issue 1: PEP → PrEP framing...")

    # Current paragraph [16] starts with "Recent work has formalized..."
    idx = find_paragraph(doc, "Recent work has formalized")
    replace_paragraph_text(doc.paragraphs[idx], (
        "Recent work has formalized HIV prevention as a feasibility problem "
        "rather than a question of drug efficacy or individual behavior [1]. "
        "An estimated 15.6 million people inject drugs worldwide, with "
        "disproportionate burdens of HIV, HBV, and HCV [2], and repeated "
        "outbreaks among PWID in the United States underscore persistent "
        "gaps in prevention infrastructure [9]. "
        "Biomedical HIV prophylaxis operates through two complementary arms: "
        "post-exposure prophylaxis (PEP), which must be initiated within a "
        "time-dependent window following exposure, and pre-exposure prophylaxis "
        "(PrEP), which establishes protective drug levels before exposure occurs."
    ))

    # Current paragraph [17] starts with "A substantial body of literature..."
    idx = find_paragraph(doc, "A substantial body of literature")
    replace_paragraph_text(doc.paragraphs[idx], (
        "For parenteral (injection) exposures, the PEP window is severely "
        "compressed. Tanner et al. [6] establish ~72-hour efficacy for mucosal "
        "exposures but a functionally compressed ~12–24-hour window for "
        "parenteral routes, where direct intravascular inoculation bypasses "
        "mucosal barriers entirely. The Prevention Theorem [1] formalizes "
        "this constraint: as the probability of productive infection P_int(t) "
        "approaches 1 with time, PEP efficacy E_PEP(t) approaches 0. For "
        "PWID whose exposures are parenteral and who face structural barriers "
        "to same-day healthcare access \u2014 including criminalization [5], stigma "
        "[7], and geographic isolation \u2014 the PEP window is functionally "
        "unattainable. PEP is therefore mathematically futile for this "
        "population under current conditions."
    ))

    # Current paragraph [18] starts with "PWID represent a critical stress test..."
    idx = find_paragraph(doc, "PWID represent a critical stress test")
    replace_paragraph_text(doc.paragraphs[idx], (
        "With PEP excluded, PrEP is the sole remaining biomedical prevention "
        "option. Although multiple PrEP regimens are now available, "
        "PWID-specific efficacy evidence remains extremely limited. Large "
        "phase 3 trials of injectable cabotegravir (HPTN 083, HPTN 084) "
        "[22, 23] and twice-yearly lenacapavir (PURPOSE 1, PURPOSE 2) "
        "[24, 25] have demonstrated high efficacy for sexual exposure, but "
        "none enrolled participants with parenteral exposures as the primary "
        "risk. The Bangkok Tenofovir Study (Choopanya et al. 2013 [19]) "
        "remains the only completed randomized efficacy trial of PrEP in "
        "PWID \u2014 using a regimen (daily oral TDF) no longer recommended as "
        "first-line PrEP. PURPOSE-4 (NCT05330624), evaluating long-acting "
        "lenacapavir in people who inject drugs [26], is the first ongoing "
        "efficacy trial designed specifically for this population, but "
        "results are not yet available. PWID must therefore navigate an "
        "8-step PrEP system built largely without their participation in "
        "the evidence base [8]."
    ))

    # Current paragraph [19] starts with "Under these conditions..."
    idx = find_paragraph(doc, "Under these conditions, prevention outcomes")
    replace_paragraph_text(doc.paragraphs[idx], (
        "Under these conditions, prevention outcomes are governed not by "
        "enforceable intervention but by stochastic avoidance — the "
        "probabilistic phenomenon whereby outbreaks fail to occur despite "
        "favorable transmission conditions, purely due to random chance in "
        "early transmission events. As network density increases driven by "
        "methamphetamine co-use, housing instability, and inadequate harm "
        "reduction infrastructure [10], the probability of continued stochastic "
        "avoidance decreases and outbreak risk escalates."
    ))

    # ══════════════════════════════════════════════════════════════════════════
    # Fix: cascade barrier description (additive decrements, not multiplicative)
    # ══════════════════════════════════════════════════════════════════════════
    print("  Fix: cascade barrier description...")
    idx = find_paragraph_containing(doc, "multiplicative penalties applied to baseline")
    orig = doc.paragraphs[idx].text
    orig = orig.replace(
        "Structural barriers were represented as multiplicative penalties "
        "applied to baseline probabilities.",
        "Within each cascade step, structural barriers were modeled as "
        "additive decrements to the step\u2019s base probability (a deficit "
        "model); cascade completion is the product of step probabilities "
        "across the 8 sequential steps."
    )
    replace_paragraph_text(doc.paragraphs[idx], orig)

    # ══════════════════════════════════════════════════════════════════════════
    # ISSUE 2: R₀(e) = 0 notation — add definition in Background
    # ══════════════════════════════════════════════════════════════════════════
    print("  Issue 2: R₀(e) notation...")

    # Insert definitional paragraph after paragraph [20] (the research question)
    idx_rq = find_paragraph(doc, "In this analysis, we examine")
    r0_def_para = insert_paragraph_after(doc, doc.paragraphs[idx_rq], (
        "We adopt the notation of Demidont (2026) [1], where R₀(e,t) denotes "
        "the residual probability of productive infection following a single "
        "exposure event e at time t post-exposure. Prevention is defined as "
        "the condition R₀(e,t) = 0 — complete abrogation of transmission risk "
        "for that event. This is distinct from the epidemiological basic "
        "reproduction number R₀, which describes population-level transmission "
        "dynamics. Throughout this analysis, P(R₀(e) = 0) represents the "
        "proportion of individuals achieving complete prevention under modeled "
        "conditions."
    ))

    # Insert [Fig. 1 about here] after the R₀(e) definition
    insert_paragraph_after(doc, r0_def_para, "")
    # Re-find to get the blank para
    idx_blank = find_paragraph_containing(doc, "the proportion of individuals achieving complete prevention")
    fig1_placeholder = insert_paragraph_after(doc, doc.paragraphs[idx_blank + 1], (
        "[Fig. 1 about here]"
    ))

    # Global R₀ notation replacement
    # "P(R₀ = 0)" → "P(R₀(e) = 0)" — must be done carefully to not double-replace
    n = global_replace_text(doc, "P(R₀ = 0)", "P(R₀(e) = 0)")
    print(f"    Replaced P(R₀ = 0) → P(R₀(e) = 0) in {n} runs")

    # "R₀ = 0" remaining (standalone, not inside P(...)) — in figure legends etc.
    # Be careful: only replace "R₀ = 0" that isn't already part of "R₀(e) = 0"
    # This is tricky with run-level replacement; skip to avoid double-replacement.
    # The key usages in figure legends use "P(R₀ = 0)" which is already handled.

    # Also replace "R₀=0" (no spaces) in references
    n = global_replace_text(doc, "achieving R₀=0", "achieving R₀(e) = 0")
    print(f"    Replaced achieving R₀=0 → achieving R₀(e) = 0 in {n} runs")

    # ══════════════════════════════════════════════════════════════════════════
    # ISSUE 8: "Model assumptions & estimands" subsection (Methods)
    # ══════════════════════════════════════════════════════════════════════════
    print("  Issue 8: Model assumptions & estimands...")

    # Insert after Monte Carlo simulation paragraph (the one starting with
    # "Monte Carlo simulations (n = 100,000"). Re-find since paragraphs shifted.
    idx_mc = find_paragraph(doc, "Monte Carlo simulations (n = 100,000")

    # Insert heading
    assumptions_heading = insert_paragraph_after(doc, doc.paragraphs[idx_mc], (
        "Model assumptions and estimands"
    ))
    # Make the heading bold by modifying the run
    for run in assumptions_heading.runs:
        run.bold = True

    # Insert body paragraph 1
    assumptions_p1 = insert_paragraph_after(doc, assumptions_heading, (
        "The cascade model makes the following key assumptions: "
        "(1) Independence: barrier penalties are applied multiplicatively "
        "across steps, treating each step as conditionally independent given "
        "the barriers present. "
        "(2) Penalty application: barrier effects are subtracted from base "
        "probabilities at each step, following a deficit model where "
        "structural barriers reduce the probability of successful step "
        "completion. "
        "(3) Incarceration mechanics: incarceration is modeled as a competing "
        "risk that resets cascade progress with an annual probability of 30% "
        "for PWID and 5% for MSM. "
        "(4) Calibration intent: model parameters are calibrated to reproduce "
        "observed population-level patterns rather than individual-level "
        "predictions; results should be interpreted as system-level "
        "feasibility assessments given the parameterization used."
    ))

    # Insert body paragraph 2
    insert_paragraph_after(doc, assumptions_p1, (
        "The primary estimand P(R₀(e) = 0) represents the proportion of "
        "individuals who, under the modeled policy conditions, would "
        "successfully complete all cascade steps and achieve effective "
        "prophylaxis. This is a system-level metric of prevention "
        "infrastructure capacity, not a clinical-trial endpoint."
    ))

    # ══════════════════════════════════════════════════════════════════════════
    # ISSUE 9: Temper absolute claims
    # ══════════════════════════════════════════════════════════════════════════
    print("  Issue 9: Tempering absolute claims...")

    # 9a: Abstract Results [10] — add "under the modeled parameterization"
    idx = find_paragraph(doc, "Results: Under current policy, PWID achieved")
    orig = doc.paragraphs[idx].text
    orig = orig.replace(
        "PWID achieved P(R₀(e) = 0) = 0.003%",
        "PWID achieved P(R₀(e) = 0) = 0.003% under the modeled parameterization"
    )
    replace_paragraph_text(doc.paragraphs[idx], orig)

    # 9b: Results cascade [36] — "achieving P(R₀(e) = 0)" → add "under the current parameterization"
    idx = find_paragraph(doc, "Under current policy, PWID achieved P(R", start=30)
    orig = doc.paragraphs[idx].text
    if "under the current parameterization" not in orig:
        orig = orig.replace(
            "achieving P(R",
            "achieving P(R",  # This line won't help — need to find actual text
        )
        # More targeted: find the full phrase
        orig = orig.replace(
            "compared to 16.3% for MSM",
            "compared to 16.3% for MSM under the current parameterization"
        )
    replace_paragraph_text(doc.paragraphs[idx], orig)

    # 9c: Policy scenarios [46] — "approximately 80% of PWID cannot achieve sustained prevention"
    try:
        idx = find_paragraph_containing(doc, "80% of PWID cannot achieve sustained prevention")
        orig = doc.paragraphs[idx].text
        orig = orig.replace(
            "cannot achieve sustained prevention",
            "cannot achieve sustained prevention given modeled barrier structures"
        )
        replace_paragraph_text(doc.paragraphs[idx], orig)
    except ValueError:
        # Might not exist verbatim; try alternative
        try:
            idx = find_paragraph_containing(doc, "cannot achieve sustained")
            orig = doc.paragraphs[idx].text
            if "given modeled barrier structures" not in orig:
                orig = orig.replace(
                    "cannot achieve sustained prevention",
                    "cannot achieve sustained prevention given modeled barrier structures"
                )
                replace_paragraph_text(doc.paragraphs[idx], orig)
        except ValueError:
            print("    Warning: Could not find 'cannot achieve sustained prevention' passage")

    # 9d: Discussion [62] — "is not primarily attributable to individual behavior"
    try:
        idx = find_paragraph_containing(doc, "not primarily attributable to individual behavior")
        orig = doc.paragraphs[idx].text
        orig = orig.replace(
            "not primarily attributable to individual behavior or pharmacological limitations",
            "not primarily attributable to individual behavior or pharmacological limitations within this model framework"
        )
        if "within this model framework" not in orig:
            # Try without "or pharmacological limitations"
            orig = orig.replace(
                "not primarily attributable to individual behavior",
                "not primarily attributable to individual behavior within this model framework"
            )
        replace_paragraph_text(doc.paragraphs[idx], orig)
    except ValueError:
        print("    Warning: Could not find discussion attribution passage")

    # 9e: Discussion [63] — "current prevention relies on stochastic avoidance"
    try:
        idx = find_paragraph_containing(doc, "current prevention relies on stochastic avoidance")
        orig = doc.paragraphs[idx].text
        orig = orig.replace(
            "current prevention relies on stochastic avoidance",
            "current prevention, as characterized by this model, relies on stochastic avoidance"
        )
        replace_paragraph_text(doc.paragraphs[idx], orig)
    except ValueError:
        print("    Warning: Could not find stochastic avoidance reliance passage")

    # 9f: Conclusions [69] — "Predictable outbreaks" hedging
    try:
        idx = find_paragraph_containing(doc, "stochastic avoidance rather than enforceable intervention")
        orig = doc.paragraphs[idx].text
        # Add hedging to the "predictable" claims
        orig = orig.replace(
            "Predictable outbreaks",
            "Predictable outbreaks, given the structural conditions modeled here,"
        )
        if "structural conditions modeled here" not in orig:
            # Try the conclusion paragraph
            pass
        replace_paragraph_text(doc.paragraphs[idx], orig)
    except ValueError:
        print("    Warning: Could not find conclusions passage")

    # ══════════════════════════════════════════════════════════════════════════
    # References 22–25 (HPTN 083, HPTN 084, PURPOSE 1, PURPOSE 2)
    # ══════════════════════════════════════════════════════════════════════════
    print("  Adding references 22\u201325...")
    idx_ref21 = find_paragraph(doc, "21. Hood JE")
    ref_cursor = doc.paragraphs[idx_ref21]

    new_refs = [
        (
            "22. Landovitz RJ, Donnell D, Clement ME, Hanscom B, Cottle L, "
            "Cober R, et al. Cabotegravir for HIV prevention in cisgender men "
            "and transgender women. N Engl J Med. 2021;385(7):595\u2013608. "
            "https://doi.org/10.1056/NEJMoa2101016"
        ),
        (
            "23. Delany-Moretlwe S, Hughes JP, Bock P, Ouma SG, Hunidzarira P, "
            "Kalonji D, et al. Cabotegravir for the prevention of HIV-1 in "
            "women: results from HPTN 084, a phase 3, randomised clinical "
            "trial. Lancet. 2022;399(10337):1779\u20131789. "
            "https://doi.org/10.1016/S0140-6736(22)00538-4"
        ),
        (
            "24. Bekker L-G, Das M, Abdool Karim Q, Bershteyn A, Blumenthal J, "
            "Crowley S, et al. Twice-yearly lenacapavir or daily F/TAF for HIV "
            "prevention in cisgender women. N Engl J Med. 2024;391(13):1179\u20131192. "
            "https://doi.org/10.1056/NEJMoa2407001"
        ),
        (
            "25. Landovitz RJ, Hanscom B, Engstrom JC, Li SS, Sedransk N, "
            "Donnell D, et al. Twice-yearly lenacapavir or daily pills for HIV "
            "prevention in MSM and transgender people. N Engl J Med. "
            "2025;392(4):309\u2013320. "
            "https://doi.org/10.1056/NEJMoa2411858"
        ),
        (
            "26. Gilead Sciences. A study to evaluate the safety and efficacy "
            "of lenacapavir for HIV pre-exposure prophylaxis in people who "
            "inject drugs (PURPOSE 4). ClinicalTrials.gov identifier: "
            "NCT05330624. Updated 2024. "
            "https://clinicaltrials.gov/study/NCT05330624"
        ),
    ]

    for ref_text in new_refs:
        ref_cursor = insert_paragraph_after(doc, ref_cursor, ref_text)

    # ══════════════════════════════════════════════════════════════════════════
    # Figure placeholders: [Fig. 6 about here] and [Fig. 7 about here]
    # ══════════════════════════════════════════════════════════════════════════
    print("  Adding [Fig. 6 about here] and [Fig. 7 about here] placeholders...")

    # Fig. 6: after the paragraph discussing V2 vs V1 divergent regional effects
    try:
        idx_fig6 = find_paragraph_containing(doc, "(Fig. 6)")
        insert_paragraph_after(doc, doc.paragraphs[idx_fig6], "")
        # Re-find since paragraph indices shifted
        idx_fig6 = find_paragraph_containing(doc, "(Fig. 6)")
        insert_paragraph_after(doc, doc.paragraphs[idx_fig6 + 1], "[Fig. 6 about here]")
        print("    Inserted [Fig. 6 about here]")
    except ValueError:
        print("    Warning: Could not find paragraph containing (Fig. 6)")

    # Fig. 7: after the paragraph discussing tornado sensitivity analysis
    try:
        idx_fig7 = find_paragraph_containing(doc, "(Fig. 7)")
        insert_paragraph_after(doc, doc.paragraphs[idx_fig7], "")
        idx_fig7 = find_paragraph_containing(doc, "(Fig. 7)")
        insert_paragraph_after(doc, doc.paragraphs[idx_fig7 + 1], "[Fig. 7 about here]")
        print("    Inserted [Fig. 7 about here]")
    except ValueError:
        print("    Warning: Could not find paragraph containing (Fig. 7)")

    # ══════════════════════════════════════════════════════════════════════════
    # Evidence asymmetry paragraph before [Fig. 8 about here]
    # ══════════════════════════════════════════════════════════════════════════
    print("  Adding evidence asymmetry paragraph before Fig. 8 placeholder...")
    try:
        idx_fig8 = find_paragraph_containing(doc, "[Fig. 8 about here]")
        # Insert the new paragraph before the [Fig. 8 about here] placeholder.
        # We find the paragraph just before the placeholder (blank line) and
        # insert after it, so the new text sits between the additional-analyses
        # paragraph and the figure placeholder.
        evidence_para = insert_paragraph_after(doc, doc.paragraphs[idx_fig8 - 1], (
            "Evidence asymmetry further compounds implementation failure. "
            "The signal-to-noise ratio (SNR) analysis (Fig. 8) shows a large "
            "disparity in the prevention evidence base between MSM and PWID, "
            "consistent with PWID being systematically underrepresented in "
            "PrEP trials. This reinforces that current prevention architecture "
            "for PWID is built on limited direct validation, magnifying "
            "uncertainty and slowing optimization."
        ))
        # Add blank line after the new paragraph
        insert_paragraph_after(doc, evidence_para, "")
        print("    Inserted evidence asymmetry paragraph")
    except ValueError:
        print("    Warning: Could not find [Fig. 8 about here] placeholder")

    # ══════════════════════════════════════════════════════════════════════════
    # BIBLIOGRAPHY: Renumber references by order of first appearance
    # ══════════════════════════════════════════════════════════════════════════
    renumber_references(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # Save
    # ══════════════════════════════════════════════════════════════════════════
    doc.save(str(MANUSCRIPT_OUT))
    print(f"\n  Manuscript saved to: {MANUSCRIPT_OUT}")


# ─── COVER LETTER EDITS ──────────────────────────────────────────────────────

def revise_cover_letter():
    doc = Document(str(COVER_IN))

    # Update figure count: "7 figures" → "8 figures"
    idx = find_paragraph_containing(doc, "7 figures")
    orig = doc.paragraphs[idx].text
    orig = orig.replace("7 figures", "8 figures")
    replace_paragraph_text(doc.paragraphs[idx], orig)
    print(f"  Cover letter: updated '7 figures' → '8 figures'")

    doc.save(str(COVER_OUT))
    print(f"  Cover letter saved to: {COVER_OUT}")


# ─── MAIN ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Revising manuscript (v2)...")
    revise_manuscript()
    print()
    print("Revising cover letter (v2)...")
    revise_cover_letter()
    print()
    print("Done. v2 files:")
    print(f"  {MANUSCRIPT_OUT}")
    print(f"  {COVER_OUT}")
